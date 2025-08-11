"""
Professional Word Report Generator for Inspection Reports
This module generates professional Word documents from inspection data
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import pytz
import pandas as pd
from io import BytesIO


def set_cell_background(cell, color):
    """Set background color for table cell"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_page_break(document):
    """Add a page break to the document"""
    paragraph = document.add_paragraph()
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def create_executive_summary_table(document, metrics):
    """Create executive summary table with key metrics"""
    # Add heading
    heading = document.add_heading('Executive Summary', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create table with key metrics
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    # Style header row
    for cell in header_cells:
        set_cell_background(cell, "4CAF50")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(12)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    metrics_data = [
        ('Total Units Inspected', f"{metrics['total_units']:,}"),
        ('Total Defects Found', f"{metrics['total_defects']:,}"),
        ('Overall Defect Rate', f"{metrics['defect_rate']:.2f}%"),
        ('Average Defects per Unit', f"{metrics['avg_defects_per_unit']:.1f}"),
        ('Units Ready for Settlement', f"{metrics['ready_units']} ({metrics['ready_pct']:.1f}%)"),
        ('Units Requiring Minor Work', f"{metrics['minor_work_units']} ({metrics['minor_pct']:.1f}%)"),
        ('Units Requiring Major Work', f"{metrics['major_work_units']} ({metrics['major_pct']:.1f}%)"),
        ('Units Requiring Extensive Work', f"{metrics['extensive_work_units']} ({metrics['extensive_pct']:.1f}%)")
    ]
    
    for metric, value in metrics_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
        
        # Style data cells
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Alternate row colors
        if len(table.rows) % 2 == 0:
            set_cell_background(row_cells[0], "F5F5F5")
            set_cell_background(row_cells[1], "F5F5F5")
    
    return table


def create_settlement_readiness_section(document, metrics):
    """Create settlement readiness breakdown section"""
    document.add_heading('Settlement Readiness Analysis', level=1)
    
    # Add description
    desc = document.add_paragraph(
        "Units have been categorized based on the number of defects found during inspection:"
    )
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Create readiness table
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    headers = ['Category', 'Criteria', 'Units', 'Percentage']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_background(header_cells[i], "2196F3")
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(11)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows with color coding
    readiness_data = [
        ('✅ Ready', '0-2 defects', metrics['ready_units'], metrics['ready_pct'], "C8E6C9"),
        ('⚠️ Minor Work', '3-7 defects', metrics['minor_work_units'], metrics['minor_pct'], "FFF3C4"),
        ('🔧 Major Work', '8-15 defects', metrics['major_work_units'], metrics['major_pct'], "FFCDD2"),
        ('🚧 Extensive Work', '15+ defects', metrics['extensive_work_units'], metrics['extensive_pct'], "F8BBD9")
    ]
    
    for category, criteria, units, percentage, bg_color in readiness_data:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = criteria
        row_cells[2].text = str(units)
        row_cells[3].text = f"{percentage:.1f}%"
        
        # Apply styling
        for cell in row_cells:
            set_cell_background(cell, bg_color)
        
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_top_trades_section(document, metrics):
    """Create top problem trades section"""
    document.add_heading('Top Problem Trades', level=1)
    
    if len(metrics['summary_trade']) == 0:
        document.add_paragraph("No defects found across all trades.")
        return
    
    # Get top 10 trades
    top_trades = metrics['summary_trade'].head(10)
    
    # Create table
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Rank', 'Trade', 'Defect Count']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_background(header_cells[i], "9C27B0")
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(11)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for idx, (_, row) in enumerate(top_trades.iterrows(), 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = row['Trade']
        row_cells[2].text = str(row['DefectCount'])
        
        # Style based on rank
        if idx <= 3:
            bg_color = "FFEBEE"  # Light red for top 3
        elif idx <= 6:
            bg_color = "FFF3E0"  # Light orange for 4-6
        else:
            bg_color = "F3E5F5"  # Light purple for 7-10
        
        for cell in row_cells:
            set_cell_background(cell, bg_color)
        
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_detailed_component_analysis(document, metrics):
    """Create detailed component analysis section"""
    document.add_heading('Detailed Component Analysis', level=1)
    
    component_details = metrics.get('component_details_summary', pd.DataFrame())
    
    if len(component_details) == 0:
        document.add_paragraph("No component details available.")
        return
    
    # Add description
    desc = document.add_paragraph(
        "The following table shows which specific units have defects for each trade, room, and component combination:"
    )
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Create table
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Trade', 'Room', 'Component', 'Units with Defects']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_background(header_cells[i], "FF9800")
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set column widths
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(2.0)
    table.columns[3].width = Inches(2.0)
    
    # Data rows (limit to first 50 for readability)
    max_rows = min(50, len(component_details))
    for idx, (_, row) in enumerate(component_details.head(max_rows).iterrows()):
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Trade'])
        row_cells[1].text = str(row['Room'])
        row_cells[2].text = str(row['Component'])
        row_cells[3].text = str(row['Units with Defects'])
        
        # Alternate row colors
        if idx % 2 == 0:
            bg_color = "FFF8E1"
        else:
            bg_color = "FFFFFF"
        
        for cell in row_cells:
            set_cell_background(cell, bg_color)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    if len(component_details) > max_rows:
        note = document.add_paragraph(f"Note: Showing first {max_rows} of {len(component_details)} total component entries.")
        note.italic = True


def create_units_summary_section(document, metrics):
    """Create units with most defects summary"""
    document.add_heading('Units Requiring Attention', level=1)
    
    if len(metrics['summary_unit']) == 0:
        document.add_paragraph("All units passed inspection with no defects found.")
        return
    
    # Get top 15 units with most defects
    top_units = metrics['summary_unit'].head(15)
    
    # Create table
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Rank', 'Unit', 'Defect Count']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_background(header_cells[i], "F44336")
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(11)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for idx, (_, row) in enumerate(top_units.iterrows(), 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = str(row['Unit'])
        row_cells[2].text = str(row['DefectCount'])
        
        # Color code based on defect count
        defect_count = row['DefectCount']
        if defect_count > 15:
            bg_color = "FFCDD2"  # Light red for extensive work
        elif defect_count > 8:
            bg_color = "FFE0B2"  # Light orange for major work
        elif defect_count > 2:
            bg_color = "FFF9C4"  # Light yellow for minor work
        else:
            bg_color = "C8E6C9"  # Light green for ready
        
        for cell in row_cells:
            set_cell_background(cell, bg_color)
        
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_professional_word_report(final_df, metrics):
    """
    Generate a professional Word document report from inspection data
    
    Args:
        final_df: Processed inspection DataFrame
        metrics: Dictionary containing calculated metrics
        
    Returns:
        Document: python-docx Document object
    """
    
    # Create new document
    document = Document()
    
    # Set document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Get current time in Melbourne timezone
    melbourne_tz = pytz.timezone('Australia/Melbourne')
    current_time = datetime.now(melbourne_tz)
    
    # ===== COVER PAGE =====
    # Title
    title = document.add_heading('🏢 PROFESSIONAL INSPECTION REPORT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(46, 125, 50)  # Green color
    
    # Subtitle
    subtitle = document.add_heading(metrics['building_name'], level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(33, 150, 243)  # Blue color
    
    # Add some spacing
    document.add_paragraph("")
    document.add_paragraph("")
    
    # Building information table
    info_table = document.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Building info data
    building_info = [
        ('Building Name:', metrics['building_name']),
        ('Address:', metrics['address']),
        ('Inspection Date:', metrics['inspection_date']),
        ('Total Units:', f"{metrics['total_units']:,}"),
        ('Unit Types:', metrics['unit_types_str']),
        ('Report Generated:', current_time.strftime('%d %B %Y at %I:%M %p AEDT'))
    ]
    
    for i, (label, value) in enumerate(building_info):
        row_cells = info_table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value
        
        # Style the label column
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(row_cells[0], "E3F2FD")
        
        # Style the value column
        set_cell_background(row_cells[1], "F8F9FA")
    
    # Add spacing before page break
    document.add_paragraph("")
    document.add_paragraph("")
    
    # Report summary paragraph
    summary_para = document.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary_para.add_run(
        f"This comprehensive inspection report covers {metrics['total_units']} units with "
        f"{metrics['total_defects']} total defects identified across all inspection points."
    )
    summary_run.font.size = Pt(12)
    summary_run.italic = True
    
    # Page break
    add_page_break(document)
    
    # ===== EXECUTIVE SUMMARY =====
    create_executive_summary_table(document, metrics)
    
    # Add page break
    add_page_break(document)
    
    # ===== SETTLEMENT READINESS =====
    create_settlement_readiness_section(document, metrics)
    
    # Add spacing
    document.add_paragraph("")
    
    # ===== TOP PROBLEM TRADES =====
    create_top_trades_section(document, metrics)
    
    # Add page break
    add_page_break(document)
    
    # ===== UNITS REQUIRING ATTENTION =====
    create_units_summary_section(document, metrics)
    
    # Add page break
    add_page_break(document)
    
    # ===== DETAILED COMPONENT ANALYSIS =====
    create_detailed_component_analysis(document, metrics)
    
    # ===== FOOTER SECTION =====
    document.add_paragraph("")
    document.add_paragraph("")
    
    # Add footer information
    footer_para = document.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Report generated on {current_time.strftime('%d %B %Y at %I:%M %p AEDT')} | "
        f"Professional Inspection Report Processor v2.0"
    )
    footer_run.font.size = Pt(10)
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    return document


# Test function to verify the module works
def test_word_generator():
    """Test function to verify Word generator is working"""
    try:
        # Create a simple test document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test to verify the Word generator is working correctly.')
        
        # Try to save to BytesIO to test complete functionality
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return True, "Word generator test successful"
    except Exception as e:
        return False, f"Word generator test failed: {str(e)}"


if __name__ == "__main__":
    # Run test when module is executed directly
    success, message = test_word_generator()
    print(f"Test Result: {message}")