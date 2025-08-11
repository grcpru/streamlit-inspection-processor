# Final Fixed Word Report Generator
# Fixes: body font only, left alignment, pie chart text, component analysis

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from datetime import datetime
import pandas as pd
import os
import tempfile
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
import numpy as np

def generate_professional_word_report(processed_data, metrics, images=None):
    """
    Generate final professional Word report with specific fixes
    """
    
    try:
        # Create new document
        doc = Document()
        
        # Setup FIXED document formatting
        setup_fixed_document_formatting(doc)
        
        # Professional cover page
        add_improved_cover_page(doc, metrics, images)
        
        # Executive overview
        add_executive_overview(doc, metrics)
        
        # Inspection process summary
        add_inspection_process_summary(doc, metrics)
        
        # Units requiring most attention
        add_improved_units_analysis(doc, metrics)
        
        # Most common defects analysis
        add_most_common_defects_analysis(doc, processed_data, metrics)
        
        # Data visualization with FIXED pie chart text
        add_fixed_data_visualization(doc, processed_data, metrics)
        
        # Trade-specific defect summary with complete unit lists
        add_complete_trade_specific_summary(doc, processed_data, metrics)
        
        # FIXED component-level breakdown
        add_fixed_component_breakdown(doc, processed_data, metrics)
        
        # Strategic recommendations
        add_strategic_recommendations(doc, metrics)
        
        # Professional footer
        add_professional_footer(doc, metrics)
        
        return doc
    
    except Exception as e:
        print(f"Error in generate_professional_word_report: {e}")
        return create_error_document(e, metrics)

def setup_fixed_document_formatting(doc):
    """FIXED document formatting - only body text bigger, left alignment"""
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    styles = doc.styles
    
    # Professional title style (KEEP ORIGINAL SIZE)
    if 'ProfessionalTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('ProfessionalTitle', 1)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(28)  # ORIGINAL SIZE
        title_font.bold = True
        title_font.color.rgb = RGBColor(31, 73, 125)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(20)
    
    # Section header style (KEEP ORIGINAL SIZE)
    if 'SectionHeader' not in [s.name for s in styles]:
        section_style = styles.add_style('SectionHeader', 1)
        section_font = section_style.font
        section_font.name = 'Calibri'
        section_font.size = Pt(18)  # ORIGINAL SIZE
        section_font.bold = True
        section_font.color.rgb = RGBColor(31, 73, 125)
        section_style.paragraph_format.space_before = Pt(24)
        section_style.paragraph_format.space_after = Pt(12)
    
    # Subsection style (KEEP ORIGINAL SIZE)
    if 'SubsectionHeader' not in [s.name for s in styles]:
        subsection_style = styles.add_style('SubsectionHeader', 1)
        subsection_font = subsection_style.font
        subsection_font.name = 'Calibri'
        subsection_font.size = Pt(14)  # ORIGINAL SIZE
        subsection_font.bold = True
        subsection_font.color.rgb = RGBColor(31, 73, 125)
        subsection_style.paragraph_format.space_before = Pt(18)
        subsection_style.paragraph_format.space_after = Pt(8)
    
    # Professional body text (ONLY THIS BIGGER + LEFT ALIGNMENT)
    if 'ProfessionalBody' not in [s.name for s in styles]:
        body_style = styles.add_style('ProfessionalBody', 1)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(12)  # BIGGER: was 11, now 12
        body_font.color.rgb = RGBColor(64, 64, 64)
        body_style.paragraph_format.line_spacing = 1.15
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # LEFT instead of JUSTIFY

def add_improved_cover_page(doc, metrics, images=None):
    """Cover page with normal header sizes"""
    
    try:
        # Company logo
        if images and images.get('logo') and os.path.exists(images['logo']):
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_para.add_run()
                logo_run.add_picture(images['logo'], width=Inches(3))
                doc.add_paragraph()
            except Exception:
                pass
        
        # Main title (NORMAL SIZE)
        title_para = doc.add_paragraph()
        title_para.style = 'ProfessionalTitle'
        title_run = title_para.add_run("PRE-SETTLEMENT INSPECTION REPORT")
        
        # Building name (NORMAL SIZE)
        doc.add_paragraph()
        building_para = doc.add_paragraph()
        building_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        building_run = building_para.add_run(f"{metrics.get('building_name', 'Building Name').upper()}")
        building_run.font.name = 'Calibri'
        building_run.font.size = Pt(20)  # NORMAL SIZE
        building_run.font.bold = True
        building_run.font.color.rgb = RGBColor(31, 73, 125)
        
        # Address (NORMAL SIZE)
        doc.add_paragraph()
        address_para = doc.add_paragraph()
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_run = address_para.add_run(metrics.get('address', 'Address'))
        address_run.font.name = 'Calibri'
        address_run.font.size = Pt(14)  # NORMAL SIZE
        address_run.font.color.rgb = RGBColor(89, 89, 89)
        
        # Cover image
        if images and images.get('cover') and os.path.exists(images['cover']):
            try:
                doc.add_paragraph()
                doc.add_paragraph()
                cover_para = doc.add_paragraph()
                cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cover_run = cover_para.add_run()
                cover_run.add_picture(images['cover'], width=Inches(5.5))
            except Exception:
                pass
        
        # Metrics dashboard
        add_improved_metrics_dashboard(doc, metrics)
        
        # Report details (NORMAL SIZE)
        doc.add_paragraph()
        doc.add_paragraph()
        details_para = doc.add_paragraph()
        details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        details_text = f"""Pre-Settlement Inspection Reporting and Analysis on Defects
for Residential Units

Created {datetime.now().strftime('%d %B %Y')}

📅 Inspection Date: {metrics.get('inspection_date', 'N/A')}
📊 Units Inspected: {metrics.get('total_units', 0):,}
🔍 Components Evaluated: {metrics.get('total_inspections', 0):,}"""
        
        details_run = details_para.add_run(details_text)
        details_run.font.name = 'Calibri'
        details_run.font.size = Pt(10)  # NORMAL SIZE
        details_run.font.color.rgb = RGBColor(89, 89, 89)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in cover page: {e}")

def add_improved_metrics_dashboard(doc, metrics):
    """Metrics dashboard with normal text sizes"""
    
    try:
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for i, width in enumerate([Inches(2.2), Inches(2.2), Inches(2.2)]):
            table.columns[i].width = width
        
        # Metrics data
        metrics_data = [
            ("🏠 TOTAL UNITS", f"{metrics.get('total_units', 0):,}", "Units Inspected"),
            ("🚨 TOTAL DEFECTS", f"{metrics.get('total_defects', 0):,}", f"{metrics.get('defect_rate', 0):.1f}% Rate"),
            ("✅ READY UNITS", f"{metrics.get('ready_units', 0)}", f"{metrics.get('ready_pct', 0):.1f}%"),
            ("⚠️ MINOR WORK", f"{metrics.get('minor_work_units', 0)}", f"{metrics.get('minor_pct', 0):.1f}%"),
            ("🔧 MAJOR WORK", f"{metrics.get('major_work_units', 0)}", f"{metrics.get('major_pct', 0):.1f}%"),
            ("🚧 EXTENSIVE WORK", f"{metrics.get('extensive_work_units', 0)}", f"{metrics.get('extensive_pct', 0):.1f}%")
        ]
        
        # Fill cells with NORMAL sizes
        for i, (label, value, subtitle) in enumerate(metrics_data):
            row = i // 3
            col = i % 3
            cell = table.cell(row, col)
            cell.vertical_alignment = 1
            
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Label (NORMAL SIZE)
            label_run = para.add_run(f"{label}\n")
            label_run.font.name = 'Calibri'
            label_run.font.size = Pt(9)  # NORMAL SIZE
            label_run.font.color.rgb = RGBColor(89, 89, 89)
            
            # Value (NORMAL SIZE)
            value_run = para.add_run(f"{value}\n")
            value_run.font.name = 'Calibri'
            value_run.font.size = Pt(16)  # NORMAL SIZE
            value_run.font.bold = True
            value_run.font.color.rgb = RGBColor(31, 73, 125)
            
            # Subtitle (NORMAL SIZE)
            subtitle_run = para.add_run(subtitle)
            subtitle_run.font.name = 'Calibri'
            subtitle_run.font.size = Pt(8)  # NORMAL SIZE
            subtitle_run.font.color.rgb = RGBColor(127, 127, 127)
    
    except Exception as e:
        print(f"Error in metrics dashboard: {e}")

def add_fixed_data_visualization(doc, processed_data, metrics):
    """Data visualization with FIXED pie chart text size"""
    
    try:
        header = doc.add_paragraph("COMPREHENSIVE DATA VISUALIZATION")
        header.style = 'SectionHeader'
        
        # Trade breakdown pie chart with SMALLER TEXT
        create_fixed_defects_breakdown_chart(doc, metrics)
        
        # Trade analysis horizontal bar chart
        create_improved_trade_analysis_chart(doc, metrics)
        
        # Room analysis chart
        if 'summary_room' in metrics and len(metrics['summary_room']) > 0:
            create_room_analysis_chart(doc, metrics)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in data visualization: {e}")

def create_fixed_defects_breakdown_chart(doc, metrics):
    """FIXED pie chart with smaller text like before"""
    
    try:
        if 'summary_trade' not in metrics or len(metrics['summary_trade']) == 0:
            return
        
        breakdown_header = doc.add_paragraph("Breakdown of Defects by Trade Category")
        breakdown_header.style = 'SubsectionHeader'
        
        trade_data = metrics['summary_trade'].head(8)  # Limit to 8 for readability
        
        fig, ax = plt.subplots(figsize=(10, 8))  # Same size as before
        
        # Professional color palette
        colors = ['#1f497d', '#4f81bd', '#9cbb58', '#f79646', '#c5504b', 
                 '#8064a2', '#4bacc6', '#f79646']
        
        # Create pie chart with SMALLER text (like before)
        wedges, texts, autotexts = ax.pie(trade_data['DefectCount'], 
                                        labels=trade_data['Trade'], 
                                        colors=colors[:len(trade_data)],
                                        autopct='%1.1f%%',
                                        startangle=90,
                                        textprops={'fontsize': 10})  # SMALLER: back to 10
        
        ax.set_title('Distribution of Defects by Trade Category', 
                    fontsize=16, fontweight='bold', pad=20)  # Normal title size
        
        # Make percentage text smaller like before
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
            autotext.set_fontsize(9)  # SMALLER: back to 9
        
        # Make labels normal size
        for text in texts:
            text.set_fontsize(10)  # SMALLER: back to 10
        
        plt.tight_layout()
        add_chart_to_document(doc, fig)
        plt.close()
        
        # Add summary text with LEFT alignment
        total_defects = metrics.get('total_defects', 0)
        if len(trade_data) > 0:
            top_trade = trade_data.iloc[0]
            summary_text = f"""The analysis shows {top_trade['Trade']} as the dominant defect category, accounting for {top_trade['DefectCount']} of the total {total_defects:,} defects ({top_trade['DefectCount']/total_defects*100:.1f}% of all identified issues). This concentration indicates where immediate remediation focus should be directed."""
            
            summary_para = doc.add_paragraph(summary_text)
            summary_para.style = 'ProfessionalBody'  # Uses LEFT alignment now
    
    except Exception as e:
        print(f"Error creating pie chart: {e}")

def add_fixed_component_breakdown(doc, processed_data, metrics):
    """FIXED component-level breakdown with proper aggregation"""
    
    try:
        header = doc.add_paragraph("COMPONENT-LEVEL BREAKDOWN ANALYSIS")
        header.style = 'SectionHeader'
        
        # Generate component data with PROPER aggregation
        component_data = None
        if 'component_details_summary' in metrics and len(metrics['component_details_summary']) > 0:
            component_data = metrics['component_details_summary']
        elif processed_data is not None and len(processed_data) > 0:
            component_data = generate_fixed_component_details(processed_data)
        
        if component_data is not None and len(component_data) > 0:
            # Ensure Unit Count column exists
            if 'Unit Count' not in component_data.columns:
                if 'Affected Units' in component_data.columns:
                    component_data['Unit Count'] = component_data['Affected Units'].apply(
                        lambda x: len(str(x).split(', ')) if pd.notna(x) else 0
                    )
                else:
                    component_data['Unit Count'] = 1
            
            # PROPER AGGREGATION: Group by Component and Trade only (not Room)
            component_aggregated = component_data.groupby(['Component', 'Trade']).agg({
                'Unit Count': 'sum',
                'Affected Units': lambda x: ', '.join(x.astype(str).unique()) if len(x) > 1 else x.iloc[0]
            }).reset_index()
            
            # Get top 15 components by actual unit count
            top_components = component_aggregated.nlargest(15, 'Unit Count')
            
            # Filter out components with only 1 unit if we have enough with more
            if len(top_components[top_components['Unit Count'] > 1]) >= 10:
                top_components = top_components[top_components['Unit Count'] > 1].head(10)
            
            # Most frequently affected components header
            most_freq_header = doc.add_paragraph("Most Frequently Affected Components")
            most_freq_header.style = 'SubsectionHeader'
            
            if len(top_components) > 0:
                # Create summary table
                comp_table = doc.add_table(rows=1, cols=5)
                comp_table.style = 'Table Grid'
                
                # Set column widths
                comp_table.columns[0].width = Inches(2.0)  # Component
                comp_table.columns[1].width = Inches(1.8)  # Trade
                comp_table.columns[2].width = Inches(2.5)  # Affected Units
                comp_table.columns[3].width = Inches(0.8)  # Count
                comp_table.columns[4].width = Inches(1.0)  # Percentage
                
                # Headers
                headers = ['Component', 'Trade', 'Sample Affected Units', 'Total Count', 'Percentage']
                for i, header in enumerate(headers):
                    cell = comp_table.cell(0, i)
                    cell.text = header
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)  # Normal header size
                    
                    # Set header background
                    set_cell_background_color(cell, "1F497D")
                
                # Add data
                total_units = metrics.get('total_units', 1)
                for _, comp_row in top_components.iterrows():
                    row = comp_table.add_row()
                    
                    # Component
                    row.cells[0].text = str(comp_row.get('Component', 'N/A'))
                    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                    
                    # Trade
                    row.cells[1].text = str(comp_row.get('Trade', 'N/A'))
                    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                    
                    # Sample affected units
                    affected_units = str(comp_row.get('Affected Units', ''))
                    if len(affected_units) > 30:
                        units_list = affected_units.split(', ')
                        sample_units = ', '.join(units_list[:5])
                        if len(units_list) > 5:
                            sample_units += f" (+ {len(units_list)-5} more)"
                        row.cells[2].text = sample_units
                    else:
                        row.cells[2].text = affected_units
                    row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
                    
                    # Count
                    unit_count = comp_row.get('Unit Count', 0)
                    row.cells[3].text = str(unit_count)
                    row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[3].paragraphs[0].runs[0].font.size = Pt(10)
                    row.cells[3].paragraphs[0].runs[0].font.bold = True
                    
                    # Percentage
                    percentage = (unit_count / total_units * 100) if total_units > 0 else 0
                    row.cells[4].text = f"{percentage:.1f}%"
                    row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[4].paragraphs[0].runs[0].font.size = Pt(10)
                    row.cells[4].paragraphs[0].runs[0].font.bold = True
                
                # IMPROVED analysis text
                if len(top_components) > 0:
                    top_component = top_components.iloc[0]
                    unit_count = top_component.get('Unit Count', 0)
                    component_name = top_component.get('Component', 'Unknown')
                    trade_name = top_component.get('Trade', 'Unknown')
                    
                    doc.add_paragraph()
                    
                    analysis_text = f"""The component-level analysis reveals "{component_name}" as the most frequently affected component, impacting {unit_count} units ({unit_count/total_units*100:.1f}% of all units inspected). This pattern indicates a systematic issue requiring focused attention within the {trade_name} trade category.

Key insights from the component analysis:
• The top 5 most problematic components affect {top_components.head(5)['Unit Count'].sum()} units combined
• {trade_name} trade shows the highest frequency of component defects
• Multiple units experiencing identical component failures suggests systematic installation or quality control issues

The concentration of defects in specific components indicates opportunities for targeted interventions, improved installation procedures, and enhanced quality control measures during the construction process."""
                    
                    analysis_para = doc.add_paragraph(analysis_text)
                    analysis_para.style = 'ProfessionalBody'  # Uses LEFT alignment
            else:
                # No significant patterns found
                no_pattern_text = "The component-level analysis shows distributed defects across various components without significant concentration patterns. This indicates isolated issues rather than systematic component problems."
                
                no_pattern_para = doc.add_paragraph(no_pattern_text)
                no_pattern_para.style = 'ProfessionalBody'
        
        else:
            no_data_para = doc.add_paragraph("Component-level breakdown data is not available for detailed analysis.")
            no_data_para.style = 'ProfessionalBody'
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in fixed component breakdown: {e}")

def generate_fixed_component_details(processed_data):
    """Generate component details with proper handling of duplicates"""
    
    try:
        required_columns = ['StatusClass', 'Trade', 'Room', 'Component', 'Unit']
        missing_columns = [col for col in required_columns if col not in processed_data.columns]
        
        if missing_columns:
            print(f"Missing columns: {missing_columns}")
            return pd.DataFrame()
        
        # Filter defects
        defects_only = processed_data[processed_data['StatusClass'] == 'Not OK']
        
        if len(defects_only) == 0:
            return pd.DataFrame()
        
        # Group by Trade, Room, Component and get ALL units
        component_summary = defects_only.groupby(['Trade', 'Room', 'Component']).agg({
            'Unit': lambda x: ', '.join(sorted(x.astype(str).unique()))
        }).reset_index()
        
        component_summary.columns = ['Trade', 'Room', 'Component', 'Affected Units']
        
        # Add unit count
        unit_counts = defects_only.groupby(['Trade', 'Room', 'Component'])['Unit'].nunique().reset_index()
        component_summary = component_summary.merge(unit_counts, on=['Trade', 'Room', 'Component'])
        component_summary.columns = ['Trade', 'Room', 'Component', 'Affected Units', 'Unit Count']
        
        # Sort by unit count (descending) then by trade
        component_summary = component_summary.sort_values(['Unit Count', 'Trade'], ascending=[False, True])
        
        return component_summary
    
    except Exception as e:
        print(f"Error in generate_fixed_component_details: {e}")
        return pd.DataFrame()

# Include all other functions from previous version with LEFT alignment for body text
def add_executive_overview(doc, metrics):
    """Executive overview with LEFT alignment"""
    
    try:
        header = doc.add_paragraph("EXECUTIVE OVERVIEW")
        header.style = 'SectionHeader'
        
        overview_text = f"""This comprehensive overview has been compiled from the pre-settlement inspection checklists for {metrics.get('total_units', 0):,} residential units within {metrics.get('building_name', 'the building')}, conducted on {metrics.get('inspection_date', 'the inspection date')}. This report was created {datetime.now().strftime('%d %B %Y')}.

Each inspection checklist details a systematic room-by-room evaluation of specific units, identifying defects and issues across all major building components including structural elements, finishes, fixtures, and systems. The inspections were conducted in accordance with industry standards and best practices for pre-settlement quality assessment.

The analysis reveals {metrics.get('total_defects', 0):,} individual defects across {metrics.get('total_inspections', 0):,} inspected components, resulting in an overall defect rate of {metrics.get('defect_rate', 0):.2f}%. Settlement readiness assessment indicates that {metrics.get('ready_pct', 0):.1f}% of units ({metrics.get('ready_units', 0)} units) are ready for immediate settlement.

Key findings include systematic patterns across trade categories, with certain defect types requiring focused remediation efforts. This report provides detailed breakdowns by trade, component, and unit to facilitate efficient resource allocation and timeline planning."""
        
        overview_para = doc.add_paragraph(overview_text)
        overview_para.style = 'ProfessionalBody'  # Now uses LEFT alignment
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in executive overview: {e}")

# Add all other functions with LEFT alignment...
# (I'll include the key ones here, rest stay the same but use LEFT alignment)

def add_improved_units_analysis(doc, metrics):
    """Units analysis with LEFT aligned text"""
    
    try:
        header = doc.add_paragraph("UNITS REQUIRING MOST ATTENTION")
        header.style = 'SectionHeader'
        
        if 'summary_unit' in metrics and len(metrics['summary_unit']) > 0:
            # Create charts
            create_top_units_horizontal_chart(doc, metrics)
            create_units_distribution_chart(doc, metrics)
            
            # Summary analysis with LEFT alignment
            top_unit = metrics['summary_unit'].iloc[0]
            summary_text = f"""Based on the comprehensive analysis of defect counts across all inspected units, Unit {top_unit['Unit']} requires the most attention with {top_unit['DefectCount']} identified defects.

The analysis shows:
• {len(metrics['summary_unit'][metrics['summary_unit']['DefectCount'] > 15])} units require extensive work (15+ defects each)
• {len(metrics['summary_unit'][(metrics['summary_unit']['DefectCount'] > 7) & (metrics['summary_unit']['DefectCount'] <= 15)])} units require major work (8-15 defects each)
• {len(metrics['summary_unit'][(metrics['summary_unit']['DefectCount'] > 2) & (metrics['summary_unit']['DefectCount'] <= 7)])} units require minor work (3-7 defects each)

This distribution enables prioritized resource allocation and realistic timeline planning for pre-settlement remediation activities."""
            
            summary_para = doc.add_paragraph(summary_text)
            summary_para.style = 'ProfessionalBody'  # LEFT alignment
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in units analysis: {e}")

# Include remaining functions with necessary adjustments...
# (All paragraph.style = 'ProfessionalBody' will now use LEFT alignment instead of JUSTIFY)

# Chart creation functions (unchanged)
def create_top_units_horizontal_chart(doc, metrics):
    """Create horizontal bar chart for top units"""
    
    try:
        if 'summary_unit' not in metrics or len(metrics['summary_unit']) == 0:
            return
        
        chart_title = doc.add_paragraph("Top 20 Units Requiring Immediate Attention")
        chart_title.style = 'SubsectionHeader'
        
        top_units = metrics['summary_unit'].head(20)
        
        if len(top_units) > 0:
            fig, ax = plt.subplots(figsize=(12, 10))
            
            # Color coding
            colors = []
            for count in top_units['DefectCount']:
                if count > 25:
                    colors.append('#8B0000')  # Dark red
                elif count > 15:
                    colors.append('#DC143C')  # Red
                elif count > 7:
                    colors.append('#FF8C00')  # Orange
                elif count > 2:
                    colors.append('#FFD700')  # Yellow
                else:
                    colors.append('#2E8B57')  # Green
            
            # Create chart
            y_pos = np.arange(len(top_units))
            bars = ax.barh(y_pos, top_units['DefectCount'], color=colors, alpha=0.8, 
                          edgecolor='white', linewidth=1.5)
            
            ax.set_yticks(y_pos)
            ax.set_yticklabels([f"Unit {unit}" for unit in top_units['Unit']], fontsize=11)
            ax.set_xlabel('Number of Defects', fontsize=14, fontweight='bold')
            ax.set_title('Units Ranked by Defect Count (Highest to Lowest)',
                        fontsize=16, fontweight='bold', pad=20)
            ax.grid(axis='x', alpha=0.3, linestyle='--')
            
            # Add value labels
            for i, (bar, value) in enumerate(zip(bars, top_units['DefectCount'])):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                       f'{value}', va='center', fontweight='bold', fontsize=11)
            
            # Add legend
            from matplotlib.patches import Patch
            legend_elements = [
                Patch(facecolor='#8B0000', label='Critical (25+ defects)'),
                Patch(facecolor='#DC143C', label='Extensive (15-24 defects)'),
                Patch(facecolor='#FF8C00', label='Major (8-14 defects)'),
                Patch(facecolor='#FFD700', label='Minor (3-7 defects)'),
                Patch(facecolor='#2E8B57', label='Ready (0-2 defects)')
            ]
            ax.legend(handles=legend_elements, loc='upper right', fontsize=10)
            
            plt.tight_layout()
            add_chart_to_document(doc, fig)
            plt.close()
    
    except Exception as e:
        print(f"Error creating units chart: {e}")

def create_units_distribution_chart(doc, metrics):
    """Create units distribution chart"""
    
    try:
        chart_title = doc.add_paragraph("Distribution of Units by Defect Severity")
        chart_title.style = 'SubsectionHeader'
        
        if 'summary_unit' in metrics and len(metrics['summary_unit']) > 0:
            fig, ax = plt.subplots(figsize=(10, 6))
            
            units_data = metrics['summary_unit']
            
            categories = []
            counts = []
            colors = []
            
            # Critical (25+)
            critical_count = len(units_data[units_data['DefectCount'] > 25])
            if critical_count > 0:
                categories.append('Critical\n(25+ defects)')
                counts.append(critical_count)
                colors.append('#8B0000')
            
            # Extensive (15-24)
            extensive_count = len(units_data[(units_data['DefectCount'] >= 15) & (units_data['DefectCount'] <= 25)])
            categories.append('Extensive\n(15-24 defects)')
            counts.append(extensive_count)
            colors.append('#DC143C')
            
            # Major (8-14)
            major_count = len(units_data[(units_data['DefectCount'] >= 8) & (units_data['DefectCount'] <= 14)])
            categories.append('Major\n(8-14 defects)')
            counts.append(major_count)
            colors.append('#FF8C00')
            
            # Minor (3-7)
            minor_count = len(units_data[(units_data['DefectCount'] >= 3) & (units_data['DefectCount'] <= 7)])
            categories.append('Minor\n(3-7 defects)')
            counts.append(minor_count)
            colors.append('#FFD700')
            
            # Ready (0-2)
            ready_count = len(units_data[units_data['DefectCount'] <= 2])
            categories.append('Ready\n(0-2 defects)')
            counts.append(ready_count)
            colors.append('#2E8B57')
            
            # Create bar chart
            bars = ax.bar(categories, counts, color=colors, alpha=0.8, edgecolor='white', linewidth=2)
            
            ax.set_ylabel('Number of Units', fontsize=12, fontweight='bold')
            ax.set_title('Unit Distribution by Defect Severity Level', fontsize=16, fontweight='bold', pad=20)
            ax.grid(axis='y', alpha=0.3, linestyle='--')
            
            # Add value labels
            for bar, value in zip(bars, counts):
                if value > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.2,
                           f'{value}', ha='center', va='bottom', fontweight='bold', fontsize=12)
            
            plt.xticks(rotation=0, fontsize=10)
            plt.tight_layout()
            add_chart_to_document(doc, fig)
            plt.close()
    
    except Exception as e:
        print(f"Error creating distribution chart: {e}")

def create_improved_trade_analysis_chart(doc, metrics):
    """Trade analysis horizontal bar chart"""
    
    try:
        trade_header = doc.add_paragraph("Trade Category Defect Analysis")
        trade_header.style = 'SubsectionHeader'
        
        if 'summary_trade' not in metrics or len(metrics['summary_trade']) == 0:
            return
        
        top_trades = metrics['summary_trade'].head(10)
        
        fig, ax = plt.subplots(figsize=(12, 8))
        
        # Create horizontal bar chart
        y_pos = np.arange(len(top_trades))
        colors = plt.cm.Blues_r(np.linspace(0.3, 0.8, len(top_trades)))
        
        bars = ax.barh(y_pos, top_trades['DefectCount'], 
                      color=colors, alpha=0.8, edgecolor='white', linewidth=1.5)
        
        ax.set_yticks(y_pos)
        ax.set_yticklabels(top_trades['Trade'], fontsize=12)
        ax.set_xlabel('Number of Defects', fontsize=14, fontweight='bold')
        ax.set_title('Defects by Trade Category (Ranked by Frequency)', 
                    fontsize=16, fontweight='bold', pad=25)
        ax.grid(axis='x', alpha=0.3, linestyle='--')
        
        # Add value labels with percentage
        total_defects = metrics.get('total_defects', 1)
        for i, (bar, value) in enumerate(zip(bars, top_trades['DefectCount'])):
            percentage = (value / total_defects * 100) if total_defects > 0 else 0
            ax.text(bar.get_width() + max(top_trades['DefectCount']) * 0.02, 
                   bar.get_y() + bar.get_height()/2,
                   f'{value} ({percentage:.1f}%)', va='center', fontweight='bold', fontsize=11)
        
        plt.tight_layout()
        add_chart_to_document(doc, fig)
        plt.close()
    
    except Exception as e:
        print(f"Error creating trade chart: {e}")

def create_room_analysis_chart(doc, metrics):
    """Room analysis chart"""
    
    try:
        room_header = doc.add_paragraph("Problem Areas by Room Type")
        room_header.style = 'SubsectionHeader'
        
        top_rooms = metrics['summary_room'].head(10)
        
        fig, ax = plt.subplots(figsize=(12, 6))
        
        y_pos = np.arange(len(top_rooms))
        bars = ax.barh(y_pos, top_rooms['DefectCount'], 
                      color="#EF8234", alpha=0.7, edgecolor='white', linewidth=1)
        
        ax.set_yticks(y_pos)
        ax.set_yticklabels(top_rooms['Room'], fontsize=12)
        ax.set_xlabel('Number of Defects', fontsize=14, fontweight='bold')
        ax.set_title('Defects by Room Type', fontsize=16, fontweight='bold', pad=25)
        ax.grid(axis='x', alpha=0.3, linestyle='--')
        
        # Add value labels
        for i, (bar, value) in enumerate(zip(bars, top_rooms['DefectCount'])):
            ax.text(bar.get_width() + max(top_rooms['DefectCount']) * 0.02, 
                   bar.get_y() + bar.get_height()/2,
                   f'{value}', va='center', fontweight='bold', fontsize=11)
        
        plt.tight_layout()
        add_chart_to_document(doc, fig)
        plt.close()
    
    except Exception as e:
        print(f"Error creating room chart: {e}")

def add_inspection_process_summary(doc, metrics):
    """Inspection process with LEFT alignment"""
    
    try:
        header = doc.add_paragraph("INSPECTION PROCESS & METHODOLOGY")
        header.style = 'SectionHeader'
        
        process_text = f"""INSPECTION SCOPE & STANDARDS

The pre-settlement inspection process was conducted systematically across all {metrics.get('total_units', 0):,} residential units, evaluating {metrics.get('total_inspections', 0):,} individual components and systems. Each inspection follows a standardized checklist covering:

• Structural elements and building envelope
• Electrical systems and fixtures
• Plumbing and water systems
• Flooring, wall, and ceiling finishes
• Doors, windows, and hardware
• Kitchen and bathroom fixtures
• Built-in storage and joinery
• Balcony and outdoor areas

QUALITY ASSESSMENT CRITERIA

Items are classified using a systematic approach:
✅ Compliant: Meets standards, no action required
❌ Defect: Requires remediation before settlement
⚪ Not Applicable: Component not present or accessible

SETTLEMENT READINESS CLASSIFICATION

Units are categorized based on defect count and severity:
• Ready for Settlement (0-2 defects): Immediate settlement possible
• Minor Work Required (3-7 defects): 1-3 days remediation
• Major Work Required (8-15 defects): 1-2 weeks remediation
• Extensive Work Required (15+ defects): 2-4 weeks remediation"""
        
        process_para = doc.add_paragraph(process_text)
        process_para.style = 'ProfessionalBody'  # LEFT alignment
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in inspection process: {e}")

def add_most_common_defects_analysis(doc, processed_data, metrics):
    """Most common defects with LEFT alignment"""
    
    try:
        header = doc.add_paragraph("MOST COMMON DEFECTS ANALYSIS")
        header.style = 'SectionHeader'
        
        if 'summary_trade' in metrics and len(metrics['summary_trade']) > 0:
            top_trade = metrics['summary_trade'].iloc[0]
            
            most_common_text = f"""Based on the comprehensive analysis of {metrics.get('total_defects', 0):,} individually identified defects across all inspection reports, "{top_trade['Trade']}" represents the most frequently reported defect category, with {top_trade['DefectCount']} instances identified.

This category encompasses various specific issues within the {top_trade['Trade'].lower()} trade, including installation problems, finish quality issues, functional defects, and compliance concerns. The concentration of defects in this trade category indicates systematic issues that require immediate attention and enhanced quality control measures."""
            
            most_common_para = doc.add_paragraph(most_common_text)
            most_common_para.style = 'ProfessionalBody'  # LEFT alignment
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in defects analysis: {e}")

def add_complete_trade_specific_summary(doc, processed_data, metrics):
    """Trade summary with complete unit lists and LEFT alignment"""
    
    try:
        header = doc.add_paragraph("TRADE SPECIFIC DEFECT SUMMARY")
        header.style = 'SectionHeader'
        
        overview_text = """This section presents a consolidated list of identified defects, grouped by trade category, along with the corresponding units where each issue was observed. All affected units are listed in full to facilitate targeted resource allocation and trade-specific remediation planning."""
        
        overview_para = doc.add_paragraph(overview_text)
        overview_para.style = 'ProfessionalBody'  # LEFT alignment
        
        # Generate complete component details
        if processed_data is not None and len(processed_data) > 0:
            try:
                component_details = generate_complete_component_details(processed_data)
                add_complete_trade_tables(doc, component_details)
            except Exception as e:
                print(f"Error generating trade tables: {e}")
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in trade summary: {e}")

def generate_complete_component_details(processed_data):
    """Generate component details with no truncation"""
    
    try:
        required_columns = ['StatusClass', 'Trade', 'Room', 'Component', 'Unit']
        missing_columns = [col for col in required_columns if col not in processed_data.columns]
        
        if missing_columns:
            print(f"Missing columns: {missing_columns}")
            return pd.DataFrame()
        
        # Filter defects
        defects_only = processed_data[processed_data['StatusClass'] == 'Not OK']
        
        if len(defects_only) == 0:
            return pd.DataFrame()
        
        # Group by Trade, Room, Component and get ALL units
        component_summary = defects_only.groupby(['Trade', 'Room', 'Component']).agg({
            'Unit': lambda x: ', '.join(sorted(x.astype(str).unique()))
        }).reset_index()
        
        component_summary.columns = ['Trade', 'Room', 'Component', 'Affected Units']
        
        # Add unit count
        unit_counts = defects_only.groupby(['Trade', 'Room', 'Component'])['Unit'].nunique().reset_index()
        component_summary = component_summary.merge(unit_counts, on=['Trade', 'Room', 'Component'])
        component_summary.columns = ['Trade', 'Room', 'Component', 'Affected Units', 'Unit Count']
        
        # Sort by trade and unit count
        component_summary = component_summary.sort_values(['Trade', 'Unit Count'], ascending=[True, False])
        
        return component_summary
    
    except Exception as e:
        print(f"Error generating component details: {e}")
        return pd.DataFrame()

def add_complete_trade_tables(doc, component_details):
    """Add trade tables with COMPLETE unit lists"""
    
    try:
        if len(component_details) == 0:
            return
        
        trades = component_details['Trade'].unique()
        
        for trade in trades:
            try:
                trade_data = component_details[component_details['Trade'] == trade]
                
                # Trade header
                trade_header = doc.add_paragraph(f"{trade}")
                trade_header.style = 'SubsectionHeader'
                
                # Create table with better widths
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Set optimal column widths for complete unit lists
                table.columns[0].width = Inches(2.5)   # Component & Location
                table.columns[1].width = Inches(4.0)   # Affected Units (WIDER)
                table.columns[2].width = Inches(0.8)   # Count
                
                # Headers
                headers = ['Component & Location', 'Affected Units', 'Count']
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)  # Normal size
                    
                    # Set header background
                    set_cell_background_color(cell, "1F497D")
                
                # Add ALL data rows (no limiting)
                for _, row in trade_data.iterrows():
                    table_row = table.add_row()
                    
                    # Component & Location
                    component_location = str(row['Component'])
                    if pd.notna(row['Room']) and str(row['Room']).strip():
                        component_location += f" ({row['Room']})"
                    
                    table_row.cells[0].text = component_location
                    table_row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                    
                    # Affected Units - COMPLETE LIST (NO TRUNCATION)
                    units_text = str(row['Affected Units'])
                    table_row.cells[1].text = units_text  # COMPLETE LIST
                    table_row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                    
                    # Count
                    table_row.cells[2].text = str(row['Unit Count'])
                    table_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table_row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
                    table_row.cells[2].paragraphs[0].runs[0].font.bold = True
                
                # Add space after each trade table
                doc.add_paragraph()
            
            except Exception as e:
                print(f"Error processing trade {trade}: {e}")
                continue
    
    except Exception as e:
        print(f"Error in trade tables: {e}")

def add_strategic_recommendations(doc, metrics):
    """Strategic recommendations with LEFT alignment"""
    
    try:
        header = doc.add_paragraph("STRATEGIC RECOMMENDATIONS & ACTION PLAN")
        header.style = 'SectionHeader'
        
        # Immediate priorities
        priorities_header = doc.add_paragraph("Immediate Priorities (Next 14 Days)")
        priorities_header.style = 'SubsectionHeader'
        
        priorities = []
        
        ready_pct = metrics.get('ready_pct', 0)
        total_defects = metrics.get('total_defects', 0)
        extensive_units = metrics.get('extensive_work_units', 0)
        
        # Settlement strategy
        if ready_pct > 75:
            priorities.append("🎯 **Fast-Track Settlement**: With 75%+ units ready, implement immediate settlement for ready units while establishing parallel remediation workflows.")
        elif ready_pct > 50:
            priorities.append("⚖️ **Phased Settlement**: Establish structured settlement phases prioritizing ready units first.")
        else:
            priorities.append("🔧 **Quality Focus**: Implement systematic remediation program before settlement to ensure customer satisfaction.")
        
        # Trade-specific priorities
        if 'summary_trade' in metrics and len(metrics['summary_trade']) > 0:
            top_trade = metrics['summary_trade'].iloc[0]
            top_trade_pct = (top_trade['DefectCount'] / total_defects * 100) if total_defects > 0 else 0
            priorities.append(f"🔧 **{top_trade['Trade']} Priority**: This trade represents {top_trade_pct:.1f}% of all defects ({top_trade['DefectCount']} instances). Assign dedicated supervision and additional resources immediately.")
        
        # Resource allocation
        if extensive_units > 0:
            priorities.append(f"👥 **Dedicated Teams**: {extensive_units} units require extensive work (15+ defects each). Establish specialized remediation teams to maintain project timeline integrity.")
        
        # Quality control
        priorities.append("📋 **Enhanced QC**: Implement additional inspection checkpoints and supervisor sign-offs for critical trades before final handover.")
        
        for i, priority in enumerate(priorities, 1):
            priority_para = doc.add_paragraph(f"{i}. {priority}")
            priority_para.style = 'ProfessionalBody'  # LEFT alignment
            priority_para.paragraph_format.left_indent = Inches(0.3)
        
        # Medium-term strategies
        doc.add_paragraph()
        medium_header = doc.add_paragraph("Medium-Term Strategies (30-60 Days)")
        medium_header.style = 'SubsectionHeader'
        
        medium_strategies = [
            "📊 **Performance Monitoring**: Establish weekly progress reviews with KPI tracking for defect remediation rates and settlement readiness improvements.",
            "🎓 **Process Improvement**: Implement lessons learned from this inspection to enhance quality control procedures for future construction phases.",
            "📞 **Stakeholder Communication**: Maintain regular communication with unit owners including progress updates and revised settlement schedules.",
            "💰 **Cost Management**: Negotiate bulk material procurement and extended trade crew availability to optimize remediation costs.",
            "🔍 **Quality Assurance**: Develop enhanced pre-handover inspection protocols to prevent similar defect patterns in future developments."
        ]
        
        for i, strategy in enumerate(medium_strategies, 1):
            strategy_para = doc.add_paragraph(f"{i}. {strategy}")
            strategy_para.style = 'ProfessionalBody'  # LEFT alignment
            strategy_para.paragraph_format.left_indent = Inches(0.3)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in recommendations: {e}")

def add_professional_footer(doc, metrics):
    """Professional footer with LEFT alignment"""
    
    try:
        header = doc.add_paragraph("REPORT INFORMATION & APPENDICES")
        header.style = 'SectionHeader'
        
        # Methodology
        methodology_header = doc.add_paragraph("Inspection Methodology & Standards")
        methodology_header.style = 'SubsectionHeader'
        
        methodology_text = """INSPECTION STANDARDS COMPLIANCE:
This comprehensive inspection was conducted in accordance with Australian building standards and industry best practices for pre-settlement quality assessment. All accessible areas and components were systematically evaluated using standardized criteria established for residential construction quality control.

CLASSIFICATION METHODOLOGY:
• ✅ Compliant: Component meets required standards, no remediation required
• ❌ Defect: Component requires attention or remediation before settlement
• ⚪ Not Applicable: Component not present, accessible, or relevant to unit type

SETTLEMENT READINESS CRITERIA:
• Ready for Settlement: 0-2 minor defects (immediate settlement possible)
• Minor Work Required: 3-7 defects (1-3 days estimated completion)
• Major Work Required: 8-15 defects (1-2 weeks estimated completion)
• Extensive Work Required: 15+ defects (2-4 weeks estimated completion)"""
        
        methodology_para = doc.add_paragraph(methodology_text)
        methodology_para.style = 'ProfessionalBody'  # LEFT alignment
        
        # Data summary
        doc.add_paragraph()
        data_summary_header = doc.add_paragraph("Comprehensive Data Summary")
        data_summary_header.style = 'SubsectionHeader'
        
        avg_defects = metrics.get('avg_defects_per_unit', 0)
        quality_score = max(0, 100 - (avg_defects * 10))
        
        data_summary_text = f"""INSPECTION METRICS OVERVIEW:
• Total Residential Units: {metrics.get('total_units', 0):,}
• Total Components Evaluated: {metrics.get('total_inspections', 0):,}
• Total Defects Identified: {metrics.get('total_defects', 0):,}
• Overall Defect Rate: {metrics.get('defect_rate', 0):.2f}%
• Average Defects per Unit: {avg_defects:.2f}
• Overall Quality Score: {quality_score:.0f}/100

SETTLEMENT READINESS DISTRIBUTION:
• Ready for Settlement: {metrics.get('ready_units', 0)} units ({metrics.get('ready_pct', 0):.1f}%)
• Minor Work Required: {metrics.get('minor_work_units', 0)} units ({metrics.get('minor_pct', 0):.1f}%)
• Major Work Required: {metrics.get('major_work_units', 0)} units ({metrics.get('major_pct', 0):.1f}%)
• Extensive Work Required: {metrics.get('extensive_work_units', 0)} units ({metrics.get('extensive_pct', 0):.1f}%)"""
        
        data_summary_para = doc.add_paragraph(data_summary_text)
        data_summary_para.style = 'ProfessionalBody'  # LEFT alignment
        
        # Report details
        doc.add_paragraph()
        details_text = f"""REPORT GENERATION DETAILS:
• Report Generated: {datetime.now().strftime('%d %B %Y at %I:%M %p')}
• Inspection Date: {metrics.get('inspection_date', 'N/A')}
• Building Complex: {metrics.get('building_name', 'N/A')}
• Property Address: {metrics.get('address', 'N/A')}

COMPANION DOCUMENTATION:
Complete defect inventories, unit-by-unit breakdowns, filterable data tables, and photographic documentation are available in the accompanying Excel workbook. This comprehensive dataset includes advanced filtering capabilities, visual dashboards, and export tools for project management systems."""
        
        details_para = doc.add_paragraph(details_text)
        details_para.style = 'ProfessionalBody'  # LEFT alignment
    
    except Exception as e:
        print(f"Error in footer: {e}")

def add_chart_to_document(doc, fig):
    """Helper function to add chart"""
    
    try:
        chart_buffer = BytesIO()
        fig.savefig(chart_buffer, format='png', dpi=300, bbox_inches='tight', 
                    facecolor='white', edgecolor='none')
        chart_buffer.seek(0)
        
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_run = chart_para.add_run()
        chart_run.add_picture(chart_buffer, width=Inches(6.5))
        
        doc.add_paragraph()
    
    except Exception as e:
        print(f"Error adding chart: {e}")

def set_cell_background_color(cell, color_hex):
    """Set cell background color"""
    
    try:
        shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    except Exception as e:
        print(f"Could not set cell background color: {e}")

def create_error_document(error, metrics):
    """Create basic document when generation fails"""
    
    doc = Document()
    title = doc.add_heading("Inspection Report - Generation Error", level=1)
    error_para = doc.add_paragraph(f"Report generation encountered an issue: {str(error)}")
    
    if metrics:
        basic_para = doc.add_paragraph(f"""
Basic Information:
Building: {metrics.get('building_name', 'N/A')}
Total Units: {metrics.get('total_units', 'N/A')}
Total Defects: {metrics.get('total_defects', 'N/A')}
        """)
    
    return doc

# Backward compatibility functions
def generate_word_report(processed_data, metrics, images=None):
    """Backward compatibility function"""
    return generate_professional_word_report(processed_data, metrics, images)

def create_inspection_report(processed_data, metrics, images=None):
    """Alternative function name"""
    return generate_professional_word_report(processed_data, metrics, images)

if __name__ == "__main__":
    print("✅ FINAL FIXED Word Report Generator loaded successfully!")
    print("\n🔧 SPECIFIC FIXES APPLIED:")
    print("• ✅ ONLY body text bigger (12pt) - headers stay normal size")
    print("• ✅ LEFT alignment for all paragraphs (no more justify)")
    print("• ✅ Pie chart text back to smaller size (9-10pt)")
    print("• ✅ FIXED component analysis - proper aggregation by component+trade")
    print("• ✅ Complete unit lists in trade tables (no ... truncation)")
    print("• ✅ Better component breakdown shows meaningful patterns")
    print("• ✅ All text alignment issues resolved")