"""
Enhanced Professional Word Report Generator for Inspection Reports
This module generates beautiful, professional Word documents with images and advanced formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
import pytz
import pandas as pd
from io import BytesIO
import os
from PIL import Image
import tempfile


def set_cell_background(cell, color):
    """Set background color for table cell"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_gradient(cell, color1, color2):
    """Set gradient background for table cell"""
    # This is a simplified gradient effect using XML
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color1)
    tcPr.append(shd)


def add_page_break(document):
    """Add a page break to the document"""
    paragraph = document.add_paragraph()
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_image_safely(paragraph, image_path, width=None):
    """Safely add image to paragraph with error handling"""
    try:
        if image_path and os.path.exists(image_path):
            # Validate image file
            try:
                with Image.open(image_path) as img:
                    # Convert to RGB if necessary
                    if img.mode in ("RGBA", "P"):
                        img = img.convert("RGB")
                    
                    # Create temporary file for processed image
                    with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
                        img.save(tmp_file.name, "JPEG", quality=85)
                        tmp_path = tmp_file.name
                
                # Add image to document
                run = paragraph.add_run()
                if width:
                    run.add_picture(tmp_path, width=width)
                else:
                    run.add_picture(tmp_path, width=Inches(6))
                
                # Clean up temporary file
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
                return True
            except Exception as e:
                print(f"Error processing image: {e}")
                return False
        return False
    except Exception as e:
        print(f"Error adding image: {e}")
        return False


def create_professional_cover_page(document, metrics, cover_image_path=None, logo_path=None):
    """Create a stunning professional cover page"""
    
    # Logo section at the top
    if logo_path:
        logo_para = document.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if add_image_safely(logo_para, logo_path, width=Inches(2)):
            logo_para.space_after = Pt(20)
    
    # Main title with enhanced styling
    title = document.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('🏢 PROFESSIONAL INSPECTION REPORT')
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    title_run.font.bold = True
    title.space_after = Pt(10)
    
    # Subtitle with building name
    subtitle = document.add_heading('', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(metrics['building_name'])
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(33, 150, 243)
    subtitle_run.font.bold = True
    subtitle.space_after = Pt(30)
    
    # Cover image if provided
    if cover_image_path:
        cover_para = document.add_paragraph()
        cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if add_image_safely(cover_para, cover_image_path, width=Inches(5)):
            cover_para.space_after = Pt(30)
    
    # Professional info box
    info_table = document.add_table(rows=7, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Enhanced building info with icons
    building_info = [
        ('🏢 Building Name', metrics['building_name']),
        ('📍 Address', metrics['address']),
        ('📅 Inspection Date', metrics['inspection_date']),
        ('🏠 Total Units', f"{metrics['total_units']:,} units"),
        ('🏗️ Unit Types', metrics['unit_types_str']),
        ('📊 Total Defects Found', f"{metrics['total_defects']:,} defects"),
        ('⏰ Report Generated', datetime.now(pytz.timezone('Australia/Melbourne')).strftime('%d %B %Y at %I:%M %p AEDT'))
    ]
    
    for i, (label, value) in enumerate(building_info):
        row_cells = info_table.rows[i].cells
        
        # Style the label cell
        label_para = row_cells[0].paragraphs[0]
        label_para.clear()
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = RGBColor(55, 71, 79)
        set_cell_background(row_cells[0], "E3F2FD")
        
        # Style the value cell
        value_para = row_cells[1].paragraphs[0]
        value_para.clear()
        value_run = value_para.add_run(str(value))
        value_run.font.size = Pt(11)
        value_run.font.color.rgb = RGBColor(33, 33, 33)
        set_cell_background(row_cells[1], "FAFAFA")
        
        # Center align both cells vertically
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Executive summary box
    document.add_paragraph().space_after = Pt(20)
    
    summary_heading = document.add_heading('📋 Executive Summary', level=2)
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    summary_para = document.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary_para.add_run(
        f"This comprehensive inspection report analyzes {metrics['total_units']} units across "
        f"{metrics['total_inspections']:,} inspection points, identifying {metrics['total_defects']} "
        f"defects with an overall defect rate of {metrics['defect_rate']:.1f}%. "
        f"{metrics['ready_units']} units ({metrics['ready_pct']:.1f}%) are ready for settlement."
    )
    summary_run.font.size = Pt(12)
    summary_run.font.italic = True
    summary_run.font.color.rgb = RGBColor(66, 66, 66)


def create_enhanced_executive_summary(document, metrics, summary_image_path=None):
    """Create an enhanced executive summary with visual elements"""
    
    # Section header with gradient effect
    heading = document.add_heading('📊 DETAILED EXECUTIVE SUMMARY', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary image if provided
    if summary_image_path:
        img_para = document.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_image_safely(img_para, summary_image_path, width=Inches(4))
        img_para.space_after = Pt(20)
    
    # Create enhanced metrics table with visual styling
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row with gradient-like effect
    header_cells = table.rows[0].cells
    headers = ['📈 Metric', '📊 Value', '🎯 Status']
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "1976D2")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Enhanced data rows with status indicators
    metrics_data = [
        ('🏠 Total Units Inspected', f"{metrics['total_units']:,}", '✅ Complete'),
        ('🔍 Total Inspection Points', f"{metrics['total_inspections']:,}", '✅ Complete'),
        ('⚠️ Total Defects Found', f"{metrics['total_defects']:,}", get_defect_status(metrics['defect_rate'])),
        ('📊 Overall Defect Rate', f"{metrics['defect_rate']:.2f}%", get_rate_status(metrics['defect_rate'])),
        ('🎯 Average Defects per Unit', f"{metrics['avg_defects_per_unit']:.1f}", get_avg_status(metrics['avg_defects_per_unit'])),
        ('✅ Units Ready for Settlement', f"{metrics['ready_units']} ({metrics['ready_pct']:.1f}%)", '🟢 Excellent'),
        ('⚠️ Units Requiring Minor Work', f"{metrics['minor_work_units']} ({metrics['minor_pct']:.1f}%)", '🟡 Attention'),
        ('🔧 Units Requiring Major Work', f"{metrics['major_work_units']} ({metrics['major_pct']:.1f}%)", '🟠 Priority'),
        ('🚧 Units Requiring Extensive Work', f"{metrics['extensive_work_units']} ({metrics['extensive_pct']:.1f}%)", '🔴 Critical')
    ]
    
    for i, (metric, value, status) in enumerate(metrics_data):
        row_cells = table.add_row().cells
        
        # Metric cell
        metric_para = row_cells[0].paragraphs[0]
        metric_para.clear()
        metric_run = metric_para.add_run(metric)
        metric_run.font.bold = True
        metric_run.font.size = Pt(10)
        
        # Value cell
        value_para = row_cells[1].paragraphs[0]
        value_para.clear()
        value_run = value_para.add_run(value)
        value_run.font.size = Pt(10)
        value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Status cell
        status_para = row_cells[2].paragraphs[0]
        status_para.clear()
        status_run = status_para.add_run(status)
        status_run.font.size = Pt(10)
        status_run.font.bold = True
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Alternate row colors
        if i % 2 == 0:
            for cell in row_cells:
                set_cell_background(cell, "F8F9FA")
        
        # Set vertical alignment
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def get_defect_status(defect_rate):
    """Get status indicator based on defect rate"""
    if defect_rate < 5:
        return "🟢 Excellent"
    elif defect_rate < 15:
        return "🟡 Good"
    elif defect_rate < 25:
        return "🟠 Fair"
    else:
        return "🔴 Poor"


def get_rate_status(rate):
    """Get status indicator based on rate"""
    if rate < 10:
        return "🟢 Low"
    elif rate < 20:
        return "🟡 Moderate"
    elif rate < 30:
        return "🟠 High"
    else:
        return "🔴 Very High"


def get_avg_status(avg):
    """Get status indicator based on average"""
    if avg < 2:
        return "🟢 Excellent"
    elif avg < 5:
        return "🟡 Good"
    elif avg < 10:
        return "🟠 Fair"
    else:
        return "🔴 Poor"


def create_visual_settlement_analysis(document, metrics, chart_image_path=None):
    """Create visually appealing settlement readiness analysis"""
    
    # Section header
    heading = document.add_heading('🏠 SETTLEMENT READINESS ANALYSIS', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add chart image if provided
    if chart_image_path:
        chart_para = document.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_image_safely(chart_para, chart_image_path, width=Inches(5))
        chart_para.space_after = Pt(20)
    
    # Add description with enhanced formatting
    desc = document.add_paragraph()
    desc_run = desc.add_run(
        "Units have been categorized into four readiness levels based on the number of defects "
        "identified during the comprehensive inspection process:"
    )
    desc_run.font.size = Pt(11)
    desc_run.font.italic = True
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    desc.space_after = Pt(15)
    
    # Create enhanced readiness table
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Enhanced header
    headers = ['🏷️ Category', '📋 Criteria', '🏠 Units', '📊 Percentage', '🎯 Priority']
    header_cells = table.rows[0].cells
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(11)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "2196F3")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Enhanced data rows with priority indicators
    readiness_data = [
        ('✅ Ready for Settlement', '0-2 defects', metrics['ready_units'], f"{metrics['ready_pct']:.1f}%", '🟢 Low', "C8E6C9"),
        ('⚠️ Minor Work Required', '3-7 defects', metrics['minor_work_units'], f"{metrics['minor_pct']:.1f}%", '🟡 Medium', "FFF3C4"),
        ('🔧 Major Work Required', '8-15 defects', metrics['major_work_units'], f"{metrics['major_pct']:.1f}%", '🟠 High', "FFCDD2"),
        ('🚧 Extensive Work Required', '15+ defects', metrics['extensive_work_units'], f"{metrics['extensive_pct']:.1f}%", '🔴 Critical', "F8BBD9")
    ]
    
    for category, criteria, units, percentage, priority, bg_color in readiness_data:
        row_cells = table.add_row().cells
        
        # Category cell
        cat_para = row_cells[0].paragraphs[0]
        cat_para.clear()
        cat_run = cat_para.add_run(category)
        cat_run.font.bold = True
        cat_run.font.size = Pt(10)
        
        # Criteria cell
        crit_para = row_cells[1].paragraphs[0]
        crit_para.clear()
        crit_run = crit_para.add_run(criteria)
        crit_run.font.size = Pt(10)
        crit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Units cell
        units_para = row_cells[2].paragraphs[0]
        units_para.clear()
        units_run = units_para.add_run(str(units))
        units_run.font.size = Pt(10)
        units_run.font.bold = True
        units_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Percentage cell
        pct_para = row_cells[3].paragraphs[0]
        pct_para.clear()
        pct_run = pct_para.add_run(percentage)
        pct_run.font.size = Pt(10)
        pct_run.font.bold = True
        pct_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Priority cell
        pri_para = row_cells[4].paragraphs[0]
        pri_para.clear()
        pri_run = pri_para.add_run(priority)
        pri_run.font.size = Pt(10)
        pri_run.font.bold = True
        pri_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply background color
        for cell in row_cells:
            set_cell_background(cell, bg_color)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_enhanced_trades_analysis(document, metrics, trades_image_path=None):
    """Create enhanced top problem trades analysis"""
    
    heading = document.add_heading('🔧 TOP PROBLEM TRADES ANALYSIS', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if len(metrics['summary_trade']) == 0:
        no_defects = document.add_paragraph()
        no_defects_run = no_defects.add_run("🎉 Excellent News! No defects found across any trades. All trades have passed inspection successfully.")
        no_defects_run.font.size = Pt(12)
        no_defects_run.font.bold = True
        no_defects_run.font.color.rgb = RGBColor(76, 175, 80)
        no_defects.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    
    # Add trades chart image if provided
    if trades_image_path:
        chart_para = document.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_image_safely(chart_para, trades_image_path, width=Inches(5))
        chart_para.space_after = Pt(20)
    
    # Get top 15 trades for more comprehensive analysis
    top_trades = metrics['summary_trade'].head(15)
    
    # Create enhanced table
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['🏆 Rank', '🔧 Trade Category', '⚠️ Defect Count', '📊 Risk Level']
    header_cells = table.rows[0].cells
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(11)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "9C27B0")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Enhanced data rows with risk assessment
    for idx, (_, row) in enumerate(top_trades.iterrows(), 1):
        row_cells = table.add_row().cells
        defect_count = row['DefectCount']
        
        # Rank cell
        rank_para = row_cells[0].paragraphs[0]
        rank_para.clear()
        rank_run = rank_para.add_run(str(idx))
        rank_run.font.size = Pt(11)
        rank_run.font.bold = True
        rank_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Trade cell
        trade_para = row_cells[1].paragraphs[0]
        trade_para.clear()
        trade_run = trade_para.add_run(row['Trade'])
        trade_run.font.size = Pt(10)
        trade_run.font.bold = True
        
        # Count cell
        count_para = row_cells[2].paragraphs[0]
        count_para.clear()
        count_run = count_para.add_run(str(defect_count))
        count_run.font.size = Pt(11)
        count_run.font.bold = True
        count_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Risk level cell
        risk_para = row_cells[3].paragraphs[0]
        risk_para.clear()
        if defect_count >= 50:
            risk_text = "🔴 Critical"
            bg_color = "FFCDD2"
        elif defect_count >= 20:
            risk_text = "🟠 High"
            bg_color = "FFE0B2"
        elif defect_count >= 10:
            risk_text = "🟡 Medium"
            bg_color = "FFF9C4"
        else:
            risk_text = "🟢 Low"
            bg_color = "C8E6C9"
        
        risk_run = risk_para.add_run(risk_text)
        risk_run.font.size = Pt(10)
        risk_run.font.bold = True
        risk_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply row styling based on rank
        for cell in row_cells:
            set_cell_background(cell, bg_color)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_recommendations_section(document, metrics):
    """Create actionable recommendations section"""
    
    heading = document.add_heading('💡 ACTIONABLE RECOMMENDATIONS', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create recommendations based on data analysis
    recommendations = []
    
    # Priority recommendations based on settlement readiness
    if metrics['extensive_work_units'] > 0:
        recommendations.append({
            'priority': '🔴 IMMEDIATE ACTION REQUIRED',
            'title': 'Critical Units Requiring Extensive Work',
            'description': f"{metrics['extensive_work_units']} units require extensive remedial work (15+ defects each). These units pose significant risk to project timeline and should be prioritized immediately.",
            'action': 'Schedule immediate contractor meetings and establish dedicated remediation teams.'
        })
    
    if metrics['major_work_units'] > 0:
        recommendations.append({
            'priority': '🟠 HIGH PRIORITY',
            'title': 'Major Work Coordination',
            'description': f"{metrics['major_work_units']} units require major remedial work (8-15 defects each). Coordinate with relevant trades to minimize delays.",
            'action': 'Develop detailed remediation schedule and assign project managers.'
        })
    
    # Trade-specific recommendations
    if len(metrics['summary_trade']) > 0:
        top_trade = metrics['summary_trade'].iloc[0]
        recommendations.append({
            'priority': '🔧 TRADE FOCUS',
            'title': f'Primary Trade Concern: {top_trade["Trade"]}',
            'description': f'The {top_trade["Trade"]} trade accounts for {top_trade["DefectCount"]} defects, representing the highest defect concentration.',
            'action': 'Conduct focused quality review with this trade contractor and implement enhanced inspection protocols.'
        })
    
    # Positive reinforcement
    if metrics['ready_pct'] > 70:
        recommendations.append({
            'priority': '🟢 POSITIVE OUTCOME',
            'title': 'Strong Overall Performance',
            'description': f"{metrics['ready_pct']:.1f}% of units are ready for settlement, indicating good overall construction quality.",
            'action': 'Document successful processes and apply best practices to remaining units.'
        })
    
    # Create recommendations table
    for i, rec in enumerate(recommendations):
        # Priority header
        priority_para = document.add_paragraph()
        priority_run = priority_para.add_run(rec['priority'])
        priority_run.font.size = Pt(12)
        priority_run.font.bold = True
        priority_para.space_before = Pt(15)
        
        # Title
        title_para = document.add_paragraph()
        title_run = title_para.add_run(rec['title'])
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(33, 33, 33)
        
        # Description
        desc_para = document.add_paragraph()
        desc_run = desc_para.add_run(rec['description'])
        desc_run.font.size = Pt(10)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Action
        action_para = document.add_paragraph()
        action_label = action_para.add_run("Recommended Action: ")
        action_label.font.bold = True
        action_label.font.size = Pt(10)
        action_text = action_para.add_run(rec['action'])
        action_text.font.size = Pt(10)
        action_text.font.italic = True


def generate_professional_word_report(final_df, metrics, images_dict=None):
    """
    Generate an enhanced professional Word document report
    
    Args:
        final_df: Processed inspection DataFrame
        metrics: Dictionary containing calculated metrics
        images_dict: Dictionary containing image paths for report enhancement
                    Keys: 'logo', 'cover', 'summary_chart', 'trades_chart', 'settlement_chart'
        
    Returns:
        Document: python-docx Document object
    """
    
    # Create new document
    document = Document()
    
    # Set document margins for professional layout
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Extract images from dictionary if provided
    images = images_dict or {}
    logo_path = images.get('logo')
    cover_image_path = images.get('cover')
    summary_chart_path = images.get('summary_chart')
    trades_chart_path = images.get('trades_chart')
    settlement_chart_path = images.get('settlement_chart')
    
    # ===== ENHANCED COVER PAGE =====
    create_professional_cover_page(document, metrics, cover_image_path, logo_path)
    
    # Page break
    add_page_break(document)
    
    # ===== ENHANCED EXECUTIVE SUMMARY =====
    create_enhanced_executive_summary(document, metrics, summary_chart_path)
    
    # Page break
    add_page_break(document)
    
    # ===== VISUAL SETTLEMENT ANALYSIS =====
    create_visual_settlement_analysis(document, metrics, settlement_chart_path)
    
    # Add spacing
    document.add_paragraph().space_after = Pt(20)
    
    # ===== ENHANCED TRADES ANALYSIS =====
    create_enhanced_trades_analysis(document, metrics, trades_chart_path)
    
    # Page break
    add_page_break(document)
    
    # ===== ACTIONABLE RECOMMENDATIONS =====
    create_recommendations_section(document, metrics)
    
    # ===== ENHANCED FOOTER =====
    document.add_paragraph().space_after = Pt(30)
    
    # Professional footer with enhanced styling
    footer_para = document.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    current_time = datetime.now(pytz.timezone('Australia/Melbourne'))
    footer_text = (f"Report generated on {current_time.strftime('%d %B %Y at %I:%M %p AEDT')} | "
                  f"Professional Inspection Report Processor v2.0 Enhanced")
    
    footer_run = footer_para.add_run(footer_text)
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Add professional disclaimer
    disclaimer_para = document.add_paragraph()
    disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer_para.add_run(
        "This report is generated based on automated analysis of inspection data. "
        "Professional judgment should be applied in conjunction with this analysis."
    )
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.italic = True
    disclaimer_run.font.color.rgb = RGBColor(150, 150, 150)
    
    return document


def create_units_priority_matrix(document, metrics):
    """Create a priority matrix for units requiring attention"""
    
    heading = document.add_heading('🎯 UNITS PRIORITY MATRIX', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if len(metrics['summary_unit']) == 0:
        success_para = document.add_paragraph()
        success_run = success_para.add_run("🎉 Outstanding Result! All units have passed inspection with no defects identified.")
        success_run.font.size = Pt(12)
        success_run.font.bold = True
        success_run.font.color.rgb = RGBColor(76, 175, 80)
        success_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    
    # Get top 20 units for comprehensive analysis
    top_units = metrics['summary_unit'].head(20)
    
    # Create enhanced table
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['🏆 Priority', '🏠 Unit ID', '⚠️ Defects', '📊 Category', '⏱️ Est. Timeline']
    header_cells = table.rows[0].cells
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(11)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "F44336")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Enhanced data rows with timeline estimates
    for idx, (_, row) in enumerate(top_units.iterrows(), 1):
        row_cells = table.add_row().cells
        defect_count = row['DefectCount']
        
        # Priority cell
        priority_para = row_cells[0].paragraphs[0]
        priority_para.clear()
        priority_run = priority_para.add_run(str(idx))
        priority_run.font.size = Pt(11)
        priority_run.font.bold = True
        priority_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Unit cell
        unit_para = row_cells[1].paragraphs[0]
        unit_para.clear()
        unit_run = unit_para.add_run(str(row['Unit']))
        unit_run.font.size = Pt(10)
        unit_run.font.bold = True
        unit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Defects cell
        defects_para = row_cells[2].paragraphs[0]
        defects_para.clear()
        defects_run = defects_para.add_run(str(defect_count))
        defects_run.font.size = Pt(11)
        defects_run.font.bold = True
        defects_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Category and timeline based on defect count
        if defect_count > 15:
            category = "🔴 Extensive"
            timeline = "3-4 weeks"
            bg_color = "FFCDD2"
        elif defect_count > 8:
            category = "🟠 Major"
            timeline = "2-3 weeks"
            bg_color = "FFE0B2"
        elif defect_count > 2:
            category = "🟡 Minor"
            timeline = "1-2 weeks"
            bg_color = "FFF9C4"
        else:
            category = "🟢 Ready"
            timeline = "< 1 week"
            bg_color = "C8E6C9"
        
        # Category cell
        cat_para = row_cells[3].paragraphs[0]
        cat_para.clear()
        cat_run = cat_para.add_run(category)
        cat_run.font.size = Pt(10)
        cat_run.font.bold = True
        cat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timeline cell
        time_para = row_cells[4].paragraphs[0]
        time_para.clear()
        time_run = time_para.add_run(timeline)
        time_run.font.size = Pt(10)
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply background color
        for cell in row_cells:
            set_cell_background(cell, bg_color)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_detailed_component_analysis(document, metrics):
    """Create enhanced detailed component analysis section"""
    
    heading = document.add_heading('🔍 DETAILED COMPONENT ANALYSIS', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    component_details = metrics.get('component_details_summary', pd.DataFrame())
    
    if len(component_details) == 0:
        no_details = document.add_paragraph()
        no_details_run = no_details.add_run("📋 No specific component defect details available for analysis.")
        no_details_run.font.size = Pt(11)
        no_details_run.font.italic = True
        no_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    
    # Add description
    desc = document.add_paragraph()
    desc_run = desc.add_run(
        "This section provides a comprehensive breakdown of defects by trade, room, and component, "
        "showing exactly which units are affected by each type of defect. This detailed analysis "
        "enables targeted remediation efforts and quality improvement initiatives."
    )
    desc_run.font.size = Pt(11)
    desc_run.font.italic = True
    desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    desc.space_after = Pt(15)
    
    # Create enhanced table
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['🔧 Trade', '🚪 Room', '⚙️ Component', '🏠 Affected Units', '📊 Impact']
    header_cells = table.rows[0].cells
    
    for i, header in enumerate(headers):
        cell = header_cells[i]
        para = cell.paragraphs[0]
        para.clear()
        run = para.add_run(header)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, "FF9800")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Set column widths
    table.columns[0].width = Inches(1.3)
    table.columns[1].width = Inches(1.2)
    table.columns[2].width = Inches(1.8)
    table.columns[3].width = Inches(2.0)
    table.columns[4].width = Inches(1.0)
    
    # Enhanced data rows (limit to first 30 for readability)
    max_rows = min(30, len(component_details))
    for idx, (_, row) in enumerate(component_details.head(max_rows).iterrows()):
        row_cells = table.add_row().cells
        
        # Trade cell
        trade_para = row_cells[0].paragraphs[0]
        trade_para.clear()
        trade_run = trade_para.add_run(str(row['Trade']))
        trade_run.font.size = Pt(9)
        trade_run.font.bold = True
        
        # Room cell
        room_para = row_cells[1].paragraphs[0]
        room_para.clear()
        room_run = room_para.add_run(str(row['Room']))
        room_run.font.size = Pt(9)
        
        # Component cell
        comp_para = row_cells[2].paragraphs[0]
        comp_para.clear()
        comp_run = comp_para.add_run(str(row['Component']))
        comp_run.font.size = Pt(9)
        
        # Units cell
        units_para = row_cells[3].paragraphs[0]
        units_para.clear()
        units_run = units_para.add_run(str(row['Units with Defects']))
        units_run.font.size = Pt(8)
        
        # Impact assessment
        unit_count = len(str(row['Units with Defects']).split(', '))
        impact_para = row_cells[4].paragraphs[0]
        impact_para.clear()
        
        if unit_count >= 10:
            impact_text = "🔴 High"
            bg_color = "FFCDD2"
        elif unit_count >= 5:
            impact_text = "🟠 Med"
            bg_color = "FFE0B2"
        elif unit_count >= 2:
            impact_text = "🟡 Low"
            bg_color = "FFF9C4"
        else:
            impact_text = "🟢 Min"
            bg_color = "C8E6C9"
        
        impact_run = impact_para.add_run(impact_text)
        impact_run.font.size = Pt(8)
        impact_run.font.bold = True
        impact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Apply styling
        for cell in row_cells:
            if idx % 2 == 0:
                set_cell_background(cell, "FFF8E1")
            else:
                set_cell_background(cell, "FFFFFF")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Override impact cell color
        set_cell_background(row_cells[4], bg_color)
    
    if len(component_details) > max_rows:
        note = document.add_paragraph()
        note_run = note.add_run(f"📝 Note: Displaying first {max_rows} of {len(component_details)} total component entries. ")
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        
        summary_run = note.add_run("Complete detailed analysis available in Excel report.")
        summary_run.font.size = Pt(9)
        summary_run.font.italic = True
        summary_run.font.color.rgb = RGBColor(100, 100, 100)


# Test function to verify the enhanced module works
def test_enhanced_word_generator():
    """Test function to verify enhanced Word generator is working"""
    try:
        # Create sample data for testing
        sample_metrics = {
            'building_name': 'Enhanced Test Building',
            'address': 'Professional Address, Melbourne VIC',
            'inspection_date': '2025-01-01',
            'unit_types_str': '2BR Apartment, 3BR Apartment',
            'total_units': 50,
            'total_inspections': 500,
            'total_defects': 75,
            'defect_rate': 15.0,
            'avg_defects_per_unit': 1.5,
            'ready_units': 35,
            'minor_work_units': 10,
            'major_work_units': 4,
            'extensive_work_units': 1,
            'ready_pct': 70.0,
            'minor_pct': 20.0,
            'major_pct': 8.0,
            'extensive_pct': 2.0,
            'summary_trade': pd.DataFrame({
                'Trade': ['Plumbing', 'Electrical', 'Painting'], 
                'DefectCount': [30, 25, 20]
            }),
            'summary_unit': pd.DataFrame({
                'Unit': ['Unit_1', 'Unit_2', 'Unit_3'], 
                'DefectCount': [15, 12, 8]
            }),
            'summary_room': pd.DataFrame({
                'Room': ['Bathroom', 'Kitchen', 'Living'], 
                'DefectCount': [25, 20, 15]
            }),
            'component_details_summary': pd.DataFrame({
                'Trade': ['Plumbing', 'Electrical'],
                'Room': ['Bathroom', 'Kitchen'],
                'Component': ['Toilet', 'GPO'],
                'Units with Defects': ['Unit_1, Unit_3', 'Unit_2, Unit_5']
            })
        }
        
        sample_data = pd.DataFrame({
            'Unit': ['Unit_1', 'Unit_2', 'Unit_3'],
            'UnitType': ['2BR Apartment', '3BR Apartment', '2BR Apartment'],
            'Room': ['Bathroom', 'Kitchen', 'Living'],
            'Component': ['Toilet', 'GPO', 'Window'],
            'StatusClass': ['Not OK', 'Not OK', 'OK'],
            'Trade': ['Plumbing', 'Electrical', 'Windows']
        })
        
        # Test images dictionary (optional)
        test_images = {
            'logo': None,  # Path to logo image
            'cover': None,  # Path to cover image
            'summary_chart': None,  # Path to summary chart
            'trades_chart': None,  # Path to trades chart
            'settlement_chart': None  # Path to settlement chart
        }
        
        # Generate enhanced document
        doc = generate_professional_word_report(sample_data, sample_metrics, test_images)
        
        # Try to save to BytesIO to test complete functionality
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return True, "Enhanced Word generator test successful with professional formatting and image support"
        
    except Exception as e:
        return False, f"Enhanced Word generator test failed: {str(e)}"


if __name__ == "__main__":
    # Run test when module is executed directly
    success, message = test_enhanced_word_generator()
    print(f"Enhanced Test Result: {message}")