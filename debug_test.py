# debug_test.py
# Simple test to debug the Word report generation issue

import streamlit as st
import sys
import traceback
from io import BytesIO
import pandas as pd
from datetime import datetime
import pytz

# Test 1: Check if python-docx is installed
st.title("🔍 Word Report Debug Test")

st.markdown("## Step 1: Check Dependencies")

try:
    from docx import Document
    st.success("✅ python-docx imported successfully")
    
    # Test creating a simple document
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    
    # Test saving to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    st.success("✅ Document creation and BytesIO conversion successful")
    
    # Show file size
    file_size = len(buffer.getvalue())
    st.info(f"Generated document size: {file_size} bytes")
    
except ImportError as e:
    st.error(f"❌ python-docx not installed: {e}")
    st.markdown("**To fix:**")
    st.code("pip install python-docx")
    st.stop()
except Exception as e:
    st.error(f"❌ Error with python-docx: {e}")
    st.code(traceback.format_exc())
    st.stop()

# Test 2: Check word_report_generator import
st.markdown("## Step 2: Check word_report_generator Module")

try:
    from word_report_generator import generate_professional_word_report
    st.success("✅ word_report_generator imported successfully")
except ImportError as e:
    st.error(f"❌ word_report_generator not found: {e}")
    st.markdown("**To fix:** Ensure word_report_generator.py is in the same directory")
    st.stop()
except Exception as e:
    st.error(f"❌ Error importing word_report_generator: {e}")
    st.code(traceback.format_exc())
    st.stop()

# Test 3: Create sample data
st.markdown("## Step 3: Create Sample Data")

# Create minimal test data
sample_data = {
    'Unit': ['Unit 1', 'Unit 2', 'Unit 3'],
    'UnitType': ['Apartment', 'Apartment', 'Townhouse'],
    'Room': ['Bedroom', 'Kitchen Area', 'Bathroom'],
    'Component': ['Walls', 'Cabinets', 'Tiles'],
    'StatusClass': ['OK', 'Not OK', 'OK'],
    'Trade': ['Painting', 'Carpentry & Joinery', 'Flooring - Tiles']
}

final_df = pd.DataFrame(sample_data)

# Create minimal metrics with ALL required keys
metrics = {
    'building_name': 'Test Building',
    'inspection_date': '2025-01-15',
    'address': '123 Test Street, Test City',
    'unit_types_str': 'Apartment, Townhouse',
    'total_units': 3,
    'total_inspections': 3,
    'total_defects': 1,
    'defect_rate': 33.3,
    'avg_defects_per_unit': 0.3,
    'ready_units': 2,
    'minor_work_units': 1,
    'major_work_units': 0,
    'extensive_work_units': 0,
    'ready_pct': 66.7,
    'minor_pct': 33.3,
    'major_pct': 0.0,
    'extensive_pct': 0.0,
    'summary_trade': pd.DataFrame({
        'Trade': ['Painting', 'Carpentry & Joinery'],
        'DefectCount': [1, 1]
    }),
    'summary_unit': pd.DataFrame({
        'Unit': ['Unit 1', 'Unit 2'],
        'DefectCount': [0, 1]
    }),
    'summary_room': pd.DataFrame({
        'Room': ['Kitchen Area'],
        'DefectCount': [1]
    }),
    'summary_unit_trade': pd.DataFrame({
        'Unit': ['Unit 2'],
        'Trade': ['Carpentry & Joinery'],
        'DefectCount': [1]
    }),
    'summary_room_comp': pd.DataFrame({
        'Room': ['Kitchen Area'],
        'Component': ['Cabinets'],
        'DefectCount': [1]
    }),
    'defects_only': pd.DataFrame({
        'Unit': ['Unit 2'],
        'UnitType': ['Apartment'],
        'Room': ['Kitchen Area'],
        'Component': ['Cabinets'],
        'StatusClass': ['Not OK'],
        'Trade': ['Carpentry & Joinery']
    }),
    'trade_specific_summary': pd.DataFrame({
        'Trade': ['Painting'],
        'Total_Defects': [1],
        'Defect_Rate_Percent': [33.3],
        'Units_Affected': [1],
        'Priority_Level': ['Low']
    }),
    'component_details_summary': pd.DataFrame({
        'Trade': ['Carpentry & Joinery'],
        'Room': ['Kitchen Area'],
        'Component': ['Cabinets'],
        'Units with Defects': ['Unit 2']
    })
}

st.success("✅ Sample data created successfully")
st.dataframe(final_df)

# Test 4: Test Word report generation
st.markdown("## Step 4: Test Word Report Generation")

if st.button("🧪 Test Generate Word Report", key="test_word_btn"):
    with st.spinner("Testing Word report generation..."):
        try:
            # Generate Word report
            word_doc = generate_professional_word_report(final_df, metrics)
            st.success("✅ Word document generated successfully!")
            
            # Convert to bytes
            word_buffer = BytesIO()
            word_doc.save(word_buffer)
            word_buffer.seek(0)
            
            # Check file size
            word_bytes = word_buffer.getvalue()
            file_size = len(word_bytes)
            st.info(f"Word document size: {file_size} bytes")
            
            if file_size > 0:
                st.success("✅ Word document conversion to bytes successful!")
                
                # Test download button
                melbourne_tz = pytz.timezone('Australia/Melbourne')
                melbourne_time = datetime.now(melbourne_tz)
                filename = f"Test_Report_{melbourne_time.strftime('%Y%m%d_%H%M%S')}.docx"
                
                st.download_button(
                    label="📄 Download Test Word Document",
                    data=word_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Download the test Word report"
                )
                
                st.success("✅ Download button created successfully!")
                st.balloons()
                
            else:
                st.error("❌ Generated document is empty!")
                
        except Exception as e:
            st.error(f"❌ Error generating Word document: {e}")
            st.code(traceback.format_exc())

# Test 5: System Information
st.markdown("## Step 5: System Information")

with st.expander("🔍 System Info"):
    st.write("**Python Version:**", sys.version)
    st.write("**Streamlit Version:**", st.__version__)
    
    # Check installed packages
    try:
        import docx
        st.write("**python-docx Version:**", docx.__version__)
    except:
        st.write("**python-docx Version:** Could not determine")
    
    try:
        import pandas as pd
        st.write("**Pandas Version:**", pd.__version__)
    except:
        st.write("**Pandas Version:** Could not determine")

# Instructions
st.markdown("## 📋 Troubleshooting Instructions")

st.markdown("""
**If any step fails, try these solutions:**

1. **Install Dependencies:**
   ```bash
   pip install python-docx pandas streamlit pytz
   ```

2. **Check File Locations:**
   - Ensure `word_report_generator.py` is in the same directory as this script
   - Both files should be in the same folder

3. **Run This Test:**
   ```bash
   streamlit run debug_test.py
   ```

4. **Common Issues:**
   - **Import Error:** Install missing packages
   - **Empty Document:** Check data structure
   - **Download Button Not Working:** Try refreshing the page
   - **Button State Issues:** Use session state properly

5. **Verify Installation:**
   ```python
   import docx
   print("python-docx installed successfully!")
   ```
""")

# Final status
st.markdown("---")
if st.session_state.get('test_completed'):
    st.success("🎉 All tests completed! Your Word report generation should work.")
else:
    st.info("👆 Click the test button above to verify Word report generation works.")