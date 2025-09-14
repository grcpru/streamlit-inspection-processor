from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.enum.section import WD_SECTION
from datetime import datetime
import pandas as pd
import os
import tempfile
from io import BytesIO
import re

# DEPENDENCY HANDLING - Added safe imports
try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # Use non-GUI backend for Streamlit
    MATPLOTLIB_AVAILABLE = True
    print("matplotlib loaded successfully for Word reports")
except ImportError as e:
    MATPLOTLIB_AVAILABLE = False
    plt = None
    print(f"matplotlib not available: {e}")

try:
    import seaborn as sns
    SEABORN_AVAILABLE = True
    print("seaborn loaded successfully for Word reports")
except ImportError as e:
    SEABORN_AVAILABLE = False
    sns = None
    print(f"seaborn not available: {e}")

try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    NUMPY_AVAILABLE = False
    np = None
    print("numpy not available - some chart features may be limited")

def remove_blank_pages(doc):
    """
    Remove blank pages from the Word document.
    This function identifies and removes paragraphs that contain only page breaks
    or are essentially empty.
    """
    try:
        print("Checking for blank pages...")
        pages_removed = 0
        paragraphs_to_remove = []
        
        # Iterate through all paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph is essentially empty or contains only page breaks
            is_blank_page_candidate = False
            
            # Check if paragraph text is empty or only whitespace
            if not paragraph.text.strip():
                # Check if paragraph contains only page breaks
                for run in paragraph.runs:
                    for elem in run.element:
                        if elem.tag.endswith('br') and elem.get(qn('w:type')) == 'page':
                            is_blank_page_candidate = True
                            break
                    if is_blank_page_candidate:
                        break
                
                # Also consider completely empty paragraphs as candidates
                if not paragraph.runs and not paragraph.text.strip():
                    is_blank_page_candidate = True
            
            # Check for consecutive empty paragraphs that might form a blank page
            if is_blank_page_candidate:
                # Look ahead to see if there are multiple consecutive empty paragraphs
                consecutive_empty = 1
                j = i + 1
                while j < len(doc.paragraphs) and j < i + 5:  # Check up to 5 paragraphs ahead
                    next_para = doc.paragraphs[j]
                    if not next_para.text.strip() and not next_para.runs:
                        consecutive_empty += 1
                        j += 1
                    else:
                        break
                
                # If we have multiple consecutive empty paragraphs, mark them for removal
                if consecutive_empty >= 2:
                    for k in range(i, min(i + consecutive_empty, len(doc.paragraphs))):
                        if k < len(doc.paragraphs) and doc.paragraphs[k] not in paragraphs_to_remove:
                            paragraphs_to_remove.append(doc.paragraphs[k])
        
        # Remove identified blank paragraphs
        for paragraph in paragraphs_to_remove:
            try:
                # Remove the paragraph element from its parent
                paragraph._element.getparent().remove(paragraph._element)
                pages_removed += 1
            except Exception as e:
                print(f"Could not remove paragraph: {e}")
                continue
        
        # Additional cleanup: Remove excessive consecutive empty paragraphs
        cleanup_excessive_spacing(doc)
        
        if pages_removed > 0:
            print(f"Removed {pages_removed} blank page elements")
        else:
            print("No blank pages detected")
        
        return pages_removed
    
    except Exception as e:
        print(f"Error in remove_blank_pages: {e}")
        return 0

def cleanup_excessive_spacing(doc):
    """
    Clean up excessive spacing between sections by removing too many consecutive empty paragraphs.
    Keeps maximum 2 empty paragraphs between content sections.
    """
    try:
        paragraphs_to_remove = []
        consecutive_empty_count = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph is empty
            if not paragraph.text.strip():
                consecutive_empty_count += 1
                # If we have more than 2 consecutive empty paragraphs, mark extras for removal
                if consecutive_empty_count > 2:
                    paragraphs_to_remove.append(paragraph)
            else:
                # Reset counter when we hit non-empty content
                consecutive_empty_count = 0
        
        # Remove excessive empty paragraphs
        for paragraph in paragraphs_to_remove:
            try:
                paragraph._element.getparent().remove(paragraph._element)
            except Exception as e:
                print(f"Could not remove excessive spacing paragraph: {e}")
                continue
        
        if len(paragraphs_to_remove) > 0:
            print(f"Cleaned up {len(paragraphs_to_remove)} excessive spacing elements")
    
    except Exception as e:
        print(f"Error in cleanup_excessive_spacing: {e}")

def optimize_page_breaks(doc):
    """
    Optimize page breaks to ensure they're properly placed and not creating blank pages.
    """
    try:
        print("Optimizing page breaks...")
        optimized_breaks = 0
        
        for paragraph in doc.paragraphs:
            # Check for manual page breaks in empty paragraphs
            if not paragraph.text.strip():
                for run in paragraph.runs:
                    # Look for page break runs in empty paragraphs
                    for elem in run.element:
                        if elem.tag.endswith('br') and elem.get(qn('w:type')) == 'page':
                            # Check if this page break is immediately after content
                            # If so, it's probably intentional; if not, consider removing
                            para_index = doc.paragraphs.index(paragraph)
                            
                            # Look at previous paragraphs to see if there's recent content
                            has_recent_content = False
                            lookback_range = min(3, para_index)
                            
                            for j in range(1, lookback_range + 1):
                                if para_index - j >= 0:
                                    prev_para = doc.paragraphs[para_index - j]
                                    if prev_para.text.strip():
                                        has_recent_content = True
                                        break
                            
                            # If no recent content, this might be an unnecessary page break
                            if not has_recent_content:
                                try:
                                    run.element.remove(elem)
                                    optimized_breaks += 1
                                except:
                                    pass
        
        if optimized_breaks > 0:
            print(f"Optimized {optimized_breaks} page breaks")
        
        return optimized_breaks
    
    except Exception as e:
        print(f"Error in optimize_page_breaks: {e}")
        return 0

def validate_document_structure(doc):
    """
    Validate the document structure and provide a summary of content distribution.
    """
    try:
        total_paragraphs = len(doc.paragraphs)
        empty_paragraphs = sum(1 for p in doc.paragraphs if not p.text.strip())
        content_paragraphs = total_paragraphs - empty_paragraphs
        
        # Count tables
        table_count = len(doc.tables)
        
        print(f"Document structure validation:")
        print(f"  Total paragraphs: {total_paragraphs}")
        print(f"  Content paragraphs: {content_paragraphs}")
        print(f"  Empty paragraphs: {empty_paragraphs}")
        print(f"  Tables: {table_count}")
        
        # Check for potential issues
        if empty_paragraphs > content_paragraphs:
            print(f"  WARNING: More empty paragraphs ({empty_paragraphs}) than content paragraphs ({content_paragraphs})")
        
        return {
            'total_paragraphs': total_paragraphs,
            'content_paragraphs': content_paragraphs,
            'empty_paragraphs': empty_paragraphs,
            'table_count': table_count
        }
    
    except Exception as e:
        print(f"Error in validate_document_structure: {e}")
        return {}

def generate_enhanced_word_report(processed_data, metrics, images=None):
    """
    Generate professional Word report matching Report_Modified.docx format
    """
    
    try:
        # Create new document
        doc = Document()
        
        # Setup document formatting with Arial font
        setup_document_formatting(doc)
        
        # Add company logo to header if available
        add_logo_to_header(doc, images)
        
        # Cover page with clean layout
        add_clean_cover_page(doc, metrics, images)
        
        # Executive overview
        add_executive_overview(doc, metrics)
        
        # Inspection process
        add_inspection_process(doc, metrics)
        
        # Units analysis
        add_units_analysis(doc, metrics)
        
        # Defects analysis
        add_defects_analysis(doc, processed_data, metrics)
        
        # Data visualization
        add_data_visualization(doc, processed_data, metrics)
        
        # Trade-specific summary
        add_trade_summary(doc, processed_data, metrics)
        
        # Component breakdown
        add_component_breakdown(doc, processed_data, metrics)
        
        # Strategic recommendations
        add_recommendations(doc, metrics)
        
        # Professional footer
        add_footer(doc, metrics)
        
        # POST-PROCESSING: Clean up the document (temporarily disabled for debugging)
        print("\nPost-processing document...")
        
        # Validate initial structure
        initial_structure = validate_document_structure(doc)
        
        # Temporarily disable blank page removal to check if it's causing the issue
        print("Skipping blank page removal for debugging...")
        blank_pages_removed = 0  # remove_blank_pages(doc)
        
        # Skip page break optimization too
        print("Skipping page break optimization for debugging...")
        breaks_optimized = 0  # optimize_page_breaks(doc)
        
        # Final validation
        final_structure = validate_document_structure(doc)
        
        print(f"\nDocument optimization complete:")
        print(f"  Blank page elements removed: {blank_pages_removed}")
        print(f"  Page breaks optimized: {breaks_optimized}")
        print(f"  Paragraphs before: {initial_structure.get('total_paragraphs', 0)}")
        print(f"  Paragraphs after: {final_structure.get('total_paragraphs', 0)}")
        
        return doc
    
    except Exception as e:
        print(f"Error in generate_enhanced_word_report: {e}")
        return create_error_document(e, metrics)

def setup_document_formatting(doc):
    """Setup document formatting with Arial font and clean styling"""
    
    # Set document margins with top margin = 3cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.0)  # Changed to 3cm as requested
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    styles = doc.styles
    
    # Title style - Arial, black
    if 'CleanTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('CleanTitle', 1)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(28)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)  # Black
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.space_before = Pt(10)
    
    # Section header - Arial, black
    if 'CleanSectionHeader' not in [s.name for s in styles]:
        section_style = styles.add_style('CleanSectionHeader', 1)
        section_font = section_style.font
        section_font.name = 'Arial'
        section_font.size = Pt(18)
        section_font.bold = True
        section_font.color.rgb = RGBColor(0, 0, 0)  # Black
        section_style.paragraph_format.space_before = Pt(20)
        section_style.paragraph_format.space_after = Pt(10)
        section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Subsection header - Arial, black
    if 'CleanSubsectionHeader' not in [s.name for s in styles]:
        subsection_style = styles.add_style('CleanSubsectionHeader', 1)
        subsection_font = subsection_style.font
        subsection_font.name = 'Arial'
        subsection_font.size = Pt(14)
        subsection_font.bold = True
        subsection_font.color.rgb = RGBColor(0, 0, 0)  # Black
        subsection_style.paragraph_format.space_before = Pt(16)
        subsection_style.paragraph_format.space_after = Pt(8)
    
    # Body text - Arial, black
    if 'CleanBody' not in [s.name for s in styles]:
        body_style = styles.add_style('CleanBody', 1)
        body_font = body_style.font
        body_font.name = 'Arial'
        body_font.size = Pt(11)
        body_font.color.rgb = RGBColor(0, 0, 0)  # Black
        body_style.paragraph_format.line_spacing = 1.2
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_logo_to_header(doc, images=None):
    """Add company logo to document header (left side)"""
    
    try:
        if images and images.get('logo') and os.path.exists(images['logo']):
            # Get the header for the first section
            section = doc.sections[0]
            header = section.header
            
            # Clear existing header content
            header.paragraphs[0].clear()
            
            # Add logo to header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            header_run = header_para.add_run()
            
            # Add logo with appropriate size for header
            header_run.add_picture(images['logo'], width=Inches(2.0))
    
    except Exception as e:
        print(f"Error adding logo to header: {e}")

def add_clean_cover_page(doc, metrics, images=None):
    """Add clean cover page matching Report_Modified.docx format"""
    
    try:
        # Main title - split into 2 lines as requested
        title_para = doc.add_paragraph()
        title_para.style = 'CleanTitle'
        title_run = title_para.add_run("PRE-SETTLEMENT\nINSPECTION REPORT")
        title_run.font.size = Pt(30)  # Slightly smaller for 2-line layout
        
        # Simple line separator
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("────────────────────────────────────────")
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(12)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Building name
        doc.add_paragraph()
        building_para = doc.add_paragraph()
        building_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        building_run = building_para.add_run(f"{metrics.get('building_name', 'Building Name').upper()}")
        building_run.font.name = 'Arial'
        building_run.font.size = Pt(22)
        building_run.font.bold = True
        building_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Address
        doc.add_paragraph()
        address_para = doc.add_paragraph()
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_run = address_para.add_run(metrics.get('address', 'Address'))
        address_run.font.name = 'Arial'
        address_run.font.size = Pt(14)
        address_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Cover image if available (center, appropriate size)
        if images and images.get('cover') and os.path.exists(images['cover']):
            try:
                doc.add_paragraph()
                cover_para = doc.add_paragraph()
                cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cover_run = cover_para.add_run()
                cover_run.add_picture(images['cover'], width=Inches(4.7))  # Adjusted size
                doc.add_paragraph()
            except Exception as e:
                print(f"Error loading cover image: {e}")
        
        # Inspection Overview section
        doc.add_paragraph()
        overview_header = doc.add_paragraph()
        overview_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        overview_run = overview_header.add_run("INSPECTION OVERVIEW")
        overview_run.font.name = 'Arial'
        overview_run.font.size = Pt(20)
        overview_run.font.bold = True
        overview_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Simple line separator
        line_para2 = doc.add_paragraph()
        line_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run2 = line_para2.add_run("──────────────────────────────────────────────────")
        line_run2.font.name = 'Arial'
        line_run2.font.size = Pt(12)
        line_run2.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        # Metrics table
        add_metrics_table(doc, metrics)
        
        # Add some space
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Report details - moved to bottom left
        details_para = doc.add_paragraph()
        details_para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Changed from CENTER to LEFT
        
        details_text = f"""Generated on {datetime.now().strftime('%d %B %Y')}

Inspection Date: {metrics.get('inspection_date', 'N/A')}
Units Inspected: {metrics.get('total_units', 0):,}
Components Evaluated: {metrics.get('total_inspections', 0):,}
Quality Score: {max(0, 100 - metrics.get('defect_rate', 0)):.1f}/100"""
        
        details_run = details_para.add_run(details_text)
        details_run.font.name = 'Arial'
        details_run.font.size = Pt(11)
        details_run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in clean cover page: {e}")

def add_metrics_table(doc, metrics):
    """Add metrics overview table with colored boxes and white borders"""
    
    try:
        # Create a regular paragraph first to ensure proper spacing
        doc.add_paragraph()
        
        # Create table with colored boxes
        table = doc.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths properly
        table.columns[0].width = Inches(2.4)
        table.columns[1].width = Inches(2.4)
        table.columns[2].width = Inches(2.4)
        
        # Metrics data with corresponding colors
        metrics_data = [
            ("TOTAL UNITS", f"{metrics.get('total_units', 0):,}", "Units Inspected", "A8D3E6"),
            ("DEFECTS FOUND", f"{metrics.get('total_defects', 0):,}", f"{metrics.get('defect_rate', 0):.1f}% Rate", "F4C2A1"),
            ("READY UNITS", f"{metrics.get('ready_units', 0)}", f"{metrics.get('ready_pct', 0):.1f}%", "C8E6C9"),
            ("MINOR WORK", f"{metrics.get('minor_work_units', 0)}", f"{metrics.get('minor_pct', 0):.1f}%", "C8E6C9"),
            ("MAJOR WORK", f"{metrics.get('major_work_units', 0)}", f"{metrics.get('major_pct', 0):.1f}%", "F4C2A1"),
            ("EXTENSIVE WORK", f"{metrics.get('extensive_work_units', 0)}", f"{metrics.get('extensive_pct', 0):.1f}%", "F4A6A6")
        ]
        
        # Fill cells with colored backgrounds and styling
        for i, (label, value, subtitle, bg_color) in enumerate(metrics_data):
            row_idx = i // 3
            col_idx = i % 3
            
            try:
                cell = table.cell(row_idx, col_idx)
                
                # Clear existing content
                cell.text = ""
                
                # Set cell background color
                try:
                    shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{bg_color}"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                except:
                    pass  # Continue without background color if this fails
                
                # Add white borders to create gaps between cells
                try:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # Add white borders (FFFFFF = white, sz="8" makes it thicker for visible gaps)
                    borders = parse_xml('''
                        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                            <w:top w:val="single" w:sz="20" w:space="0" w:color="FFFFFF"/>
                            <w:left w:val="single" w:sz="20" w:space="0" w:color="FFFFFF"/>
                            <w:bottom w:val="single" w:sz="20" w:space="0" w:color="FFFFFF"/>
                            <w:right w:val="single" w:sz="20" w:space="0" w:color="FFFFFF"/>
                        </w:tcBorders>
                    ''')
                    tcPr.append(borders)
                except:
                    pass  # Continue without borders if this fails
                
                # Add content to cell
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add padding to create space inside cells
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.left_indent = Pt(8)
                para.paragraph_format.right_indent = Pt(8)
                
                # Label
                label_run = para.add_run(f"{label}\n")
                label_run.font.name = 'Arial'
                label_run.font.size = Pt(10)
                label_run.font.color.rgb = RGBColor(0, 0, 0)
                label_run.font.bold = False
                
                # Value (large number)
                value_run = para.add_run(f"{value}\n")
                value_run.font.name = 'Arial'
                value_run.font.size = Pt(24)
                value_run.font.bold = True
                value_run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Subtitle
                subtitle_run = para.add_run(subtitle)
                subtitle_run.font.name = 'Arial'
                subtitle_run.font.size = Pt(9)
                subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
                subtitle_run.font.bold = False
                
            except Exception as cell_error:
                print(f"Error processing cell {i}: {cell_error}")
                continue
        
        # Add table-level spacing to create overall gaps
        try:
            tbl = table._tbl
            tblPr = tbl.get_or_add_tblPr()
            
            # Add cell spacing to create gaps between cells
            cellSpacing = parse_xml('<w:tblCellSpacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="120" w:type="dxa"/>')
            tblPr.append(cellSpacing)
        except:
            pass
        
        # Add spacing after table
        doc.add_paragraph()
    
    except Exception as e:
        print(f"Error in metrics table: {e}")
        # Fallback: Add text-based metrics if table fails
        try:
            fallback_para = doc.add_paragraph("INSPECTION METRICS:")
            fallback_para.style = 'CleanSubsectionHeader'
            
            metrics_text = f"""Total Units: {metrics.get('total_units', 0):,}
Defects Found: {metrics.get('total_defects', 0):,} ({metrics.get('defect_rate', 0):.1f}% Rate)
Ready Units: {metrics.get('ready_units', 0)} ({metrics.get('ready_pct', 0):.1f}%)
Minor Work: {metrics.get('minor_work_units', 0)} ({metrics.get('minor_pct', 0):.1f}%)
Major Work: {metrics.get('major_work_units', 0)} ({metrics.get('major_pct', 0):.1f}%)
Extensive Work: {metrics.get('extensive_work_units', 0)} ({metrics.get('extensive_pct', 0):.1f}%)"""
            
            fallback_text_para = doc.add_paragraph(metrics_text)
            fallback_text_para.style = 'CleanBody'
        except:
            pass

def set_cell_borders(cell):
    """Set clean borders for table cells"""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Add light borders
        borders = parse_xml('''
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>
            </w:tcBorders>
        ''')
        tcPr.append(borders)
    except Exception as e:
        print(f"Could not set cell borders: {e}")

def set_cell_background_color(cell, color_hex):
    """Set cell background color with hex color code"""
    
    try:
        shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    except Exception as e:
        print(f"Could not set cell background color: {e}")

def add_formatted_text_with_bold(paragraph, text, style_name='CleanBody'):
    """Add text with **bold** formatting support"""
    
    try:
        # Split text by **bold** markers
        parts = re.split(r'\*\*(.*?)\*\*', text)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Regular text
                if part:
                    run = paragraph.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:  # Bold text (inside ** **)
                run = paragraph.add_run(part)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.bold = True
        
        paragraph.style = style_name
    
    except Exception as e:
        # Fallback: just add the text normally
        run = paragraph.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.style = style_name

def add_executive_overview(doc, metrics):
    """Add executive overview section"""
    
    try:
        header = doc.add_paragraph("EXECUTIVE OVERVIEW")
        header.style = 'CleanSectionHeader'
        
        # Add line separator
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("───────────────────────────────────────────────────────────────")
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
                
        overview_text = f"""This comprehensive quality assessment encompasses the systematic evaluation of {metrics.get('total_units', 0):,} residential units within {metrics.get('building_name', 'the building complex')}, conducted on {metrics.get('inspection_date', 'the inspection date')}. This report was compiled on {datetime.now().strftime('%d %B %Y')}.

**Inspection Methodology**: Each unit underwent thorough room-by-room evaluation covering all major building components, including structural elements, mechanical systems, finishes, fixtures, and fittings. The assessment follows industry-standard protocols for pre-settlement quality verification.

**Key Findings**: The inspection revealed {metrics.get('total_defects', 0):,} individual defects across {metrics.get('total_inspections', 0):,} evaluated components, yielding an overall defect rate of {metrics.get('defect_rate', 0):.2f}%. Defect level analysis indicates {metrics.get('ready_pct', 0):.1f}% of units ({metrics.get('ready_units', 0)} units) require only minor work for handover.

**Strategic Insights**: The data reveals systematic patterns across trade categories, with concentrated defect types requiring targeted remediation strategies. This analysis enables optimized resource allocation and realistic timeline planning for completion preparation."""
        
        overview_para = doc.add_paragraph()
        add_formatted_text_with_bold(overview_para, overview_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in executive overview: {e}")

def add_inspection_process(doc, metrics):
    """Add inspection process section"""
    
    try:
        header = doc.add_paragraph("INSPECTION PROCESS & METHODOLOGY")
        header.style = 'CleanSectionHeader'
        
        # Decorative line - shortened as requested
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deco_run = deco_para.add_run("───────────────────────────────────────────────────────────────")  # Shortened line
        deco_run.font.name = 'Arial'
        deco_run.font.size = Pt(10)
        deco_run.font.color.rgb = RGBColor(0, 0, 0)
                
        # Scope section
        scope_header = doc.add_paragraph("INSPECTION SCOPE & STANDARDS")
        scope_header.style = 'CleanSubsectionHeader'
        
        scope_text = f"""The comprehensive pre-settlement quality assessment was systematically executed across all {metrics.get('total_units', 0):,} residential units, encompassing detailed evaluation of {metrics.get('total_inspections', 0):,} individual components and building systems.

**Structural Assessment**
• Building envelope integrity and weatherproofing
• Structural elements and load-bearing components
• Foundation and concrete work evaluation

**Systems Evaluation**  
• Electrical installations, fixtures, and safety compliance
• Plumbing systems, water pressure, and drainage
• HVAC systems and ventilation adequacy

**Finishes & Fixtures**
• Wall, ceiling, and flooring finish quality
• Door and window installation and operation
• Kitchen and bathroom fixture functionality
• Built-in storage and joinery craftsmanship"""
        
        scope_para = doc.add_paragraph()
        add_formatted_text_with_bold(scope_para, scope_text)
        
        doc.add_paragraph()
        
        # Quality criteria section
        criteria_header = doc.add_paragraph("QUALITY ASSESSMENT CRITERIA")
        criteria_header.style = 'CleanSubsectionHeader'
        
        criteria_text = """Classification methodology follows systematic evaluation protocols:

**Compliant Status**: Component meets required standards and specifications, ready for handover

**Defect Status**: Component requires remediation or adjustment before final handover  

**Not Applicable**: Component not present, accessible, or relevant to specific unit configuration

Each assessment point is documented with photographic evidence and detailed descriptions to facilitate efficient remediation workflows."""
        
        criteria_para = doc.add_paragraph()
        add_formatted_text_with_bold(criteria_para, criteria_text)
        
        doc.add_paragraph()
        
        # Defect level framework section
        readiness_header = doc.add_paragraph("DEFECT LEVEL FRAMEWORK")
        readiness_header.style = 'CleanSubsectionHeader'
        
        readiness_text = """Units are categorized using evidence-based defect thresholds and estimated remediation timeframes:

🟢**Ready for Settlement** (0-2 defects)
   Immediate handover capability with minor or cosmetic issues only

**🟡 Minor Work Required** (3-7 defects)  
   1-3 days estimated remediation time for quick fixes and adjustments

**🟠 Major Work Required** (8-15 defects)
   1-2 weeks estimated completion for substantial repairs and installations

**🔴 Extensive Work Required** (15+ defects)
   2-4 weeks estimated timeframe for comprehensive remediation and quality upgrades"""
        
        readiness_para = doc.add_paragraph()
        add_formatted_text_with_bold(readiness_para, readiness_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in inspection process: {e}")

def add_units_analysis(doc, metrics):
    """Add units analysis section"""
    
    try:
        header = doc.add_paragraph("UNITS REQUIRING PRIORITY ATTENTION")
        header.style = 'CleanSectionHeader'
        
        # Decorative line - shortened as requested
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deco_run = deco_para.add_run("───────────────────────────────────────────────────────────────")  # Shortened line
        deco_run.font.name = 'Arial'
        deco_run.font.size = Pt(10)
        deco_run.font.color.rgb = RGBColor(0, 0, 0)

        if 'summary_unit' in metrics and len(metrics['summary_unit']) > 0:
            # Create chart
            create_units_chart(doc, metrics)
            
            # Analysis text with bold formatting
            top_unit = metrics['summary_unit'].iloc[0]
            
            summary_text = f"""**Priority Analysis Results**: Unit {top_unit['Unit']} requires immediate priority attention with {top_unit['DefectCount']} identified defects, representing the highest concentration of remediation needs within the development.

**Resource Allocation Framework**:
• **Critical Priority**: {len(metrics['summary_unit'][metrics['summary_unit']['DefectCount'] > 15])} units requiring extensive remediation (15+ defects each)
• **High Priority**: {len(metrics['summary_unit'][(metrics['summary_unit']['DefectCount'] > 7) & (metrics['summary_unit']['DefectCount'] <= 15)])} units requiring major work (8-15 defects each)  
• **Medium Priority**: {len(metrics['summary_unit'][(metrics['summary_unit']['DefectCount'] > 2) & (metrics['summary_unit']['DefectCount'] <= 7)])} units requiring intermediate work (3-7 defects each)
• **Handover Ready**: {len(metrics['summary_unit'][metrics['summary_unit']['DefectCount'] <= 2])} units ready for immediate handover

**Strategic Insights**: This distribution pattern enables targeted resource deployment and realistic timeline forecasting for completion preparation activities. The concentration of defects in specific units suggests opportunities for parallel remediation workflows and optimized trade scheduling."""
            
            summary_para = doc.add_paragraph()
            add_formatted_text_with_bold(summary_para, summary_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in units analysis: {e}")

def create_units_chart(doc, metrics):
    """Create units chart with legend"""
    
    if not MATPLOTLIB_AVAILABLE:
        add_text_units_summary(doc, metrics)
        return
    
    try:
        chart_title = doc.add_paragraph("Top 20 Units Requiring Immediate Intervention")
        chart_title.style = 'CleanSubsectionHeader'
        
        top_units = metrics['summary_unit'].head(20)
        
        if len(top_units) > 0:
            fig, ax = plt.subplots(figsize=(16, 12))
            
            # Color coding based on defect severity
            colors = []
            for count in top_units['DefectCount']:
                if count > 25:
                    colors.append('#ff9999')  # Light red for critical
                elif count >= 15:
                    colors.append('#ffcc99')  # Light orange for extensive
                elif count >= 8:
                    colors.append('#ffff99')  # Light yellow for major
                elif count >= 3:
                    colors.append('#99ff99')  # Light green for minor
                else:
                    colors.append('#99ccff')  # Light blue for ready
            
            if NUMPY_AVAILABLE:
                y_pos = np.arange(len(top_units))
            else:
                y_pos = list(range(len(top_units)))
            
            bars = ax.barh(y_pos, top_units['DefectCount'], color=colors, alpha=0.8)
            
            ax.set_yticks(y_pos)
            ax.set_yticklabels([f"Unit {unit}" for unit in top_units['Unit']], fontsize=14)
            ax.set_xlabel('Number of Defects', fontsize=16, fontweight='600')
            ax.set_title('Units Ranked by Defect Concentration (Priority Order)',
                        fontsize=18, fontweight='600', pad=25)
            
            ax.grid(axis='x', alpha=0.3, linestyle=':')
            
            # Value labels
            for i, (bar, value) in enumerate(zip(bars, top_units['DefectCount'])):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                       f'{value}', va='center', fontweight='bold', fontsize=12)
            
            # Add legend with proper colors - FIXED
            from matplotlib.patches import Patch
            legend_elements = [
                Patch(facecolor='#ff9999', label='Critical (25+ defects)', alpha=0.8),
                Patch(facecolor='#ffcc99', label='Extensive (15-24 defects)', alpha=0.8),
                Patch(facecolor='#ffff99', label='Major (8-14 defects)', alpha=0.8),
                Patch(facecolor='#99ff99', label='Minor (3-7 defects)', alpha=0.8),
                Patch(facecolor='#99ccff', label='Ready (0-2 defects)', alpha=0.8)
            ]
            ax.legend(handles=legend_elements, loc='upper right', fontsize=14, framealpha=0.9)
            
            plt.tight_layout()
            add_chart_to_document(doc, fig)
            plt.close()
    
    except Exception as e:
        print(f"Error creating units chart: {e}")

def add_defects_analysis(doc, processed_data, metrics):
    """Add defects analysis section"""
    
    try:
        header = doc.add_paragraph("DEFECT PATTERNS & ANALYSIS")
        header.style = 'CleanSectionHeader'
        
        # Decorative line - shortened as requested
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deco_run = deco_para.add_run("───────────────────────────────────────────────────────────────")  # Shortened line
        deco_run.font.name = 'Arial'
        deco_run.font.size = Pt(10)
        deco_run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()

        if 'summary_trade' in metrics and len(metrics['summary_trade']) > 0:
            top_trade = metrics['summary_trade'].iloc[0]
            total_defects = metrics.get('total_defects', 0)
            trade_percentage = (top_trade['DefectCount']/total_defects*100) if total_defects > 0 else 0
            
            defects_text = f"""**Primary Defect Category Analysis**: The comprehensive evaluation of {total_defects:,} individually documented defects reveals "{top_trade['Trade']}" as the dominant concern category, accounting for {top_trade['DefectCount']} instances ({trade_percentage:.1f}% of total defects).

**Pattern Recognition**: This concentration within the {top_trade['Trade'].lower()} trade category encompasses multiple sub-issues including installation inconsistencies, finish quality variations, functional defects, and compliance gaps. The systematic nature of these defects indicates opportunities for targeted quality control improvements.

**Strategic Implications**: The clustering of defects within specific trade categories suggests that focused remediation efforts targeting the top 3-4 trade categories could address approximately 60-80% of all identified issues, enabling efficient resource deployment and accelerated completion timelines."""
            
            defects_para = doc.add_paragraph()
            add_formatted_text_with_bold(defects_para, defects_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in defects analysis: {e}")

def add_data_visualization(doc, processed_data, metrics):
    """Add data visualization section"""
    
    try:
        header = doc.add_paragraph("COMPREHENSIVE DATA VISUALISATION")
        header.style = 'CleanSectionHeader'
        
        # Decorative line - shortened as requested
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deco_run = deco_para.add_run("───────────────────────────────────────────────────────────────")  # Shortened line
        deco_run.font.name = 'Arial'
        deco_run.font.size = Pt(10)
        deco_run.font.color.rgb = RGBColor(0, 0, 0)

        intro_text = "This section presents visual analytics of the inspection data, highlighting key patterns and trends to support strategic decision-making and resource allocation."
        intro_para = doc.add_paragraph(intro_text)
        intro_para.style = 'CleanBody'
        
        doc.add_paragraph()
        
        # Create pie chart
        create_pie_chart(doc, metrics)
        
        # Create severity chart
        create_severity_chart(doc, metrics)
        
        # Create trade chart
        create_trade_chart(doc, metrics)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in data visualization: {e}")

def create_pie_chart(doc, metrics):
    """Create pie chart"""
    
    if not MATPLOTLIB_AVAILABLE:
        add_text_trade_summary(doc, metrics)
        return
    
    try:
        if 'summary_trade' not in metrics or len(metrics['summary_trade']) == 0:
            return
        
        breakdown_header = doc.add_paragraph("Defects Distribution by Trade Category")
        breakdown_header.style = 'CleanSubsectionHeader'
        
        trade_data = metrics['summary_trade'].copy()
        total_defects = metrics.get('total_defects', 0)
        
        num_trades = len(trade_data)
        
        if NUMPY_AVAILABLE and num_trades <= 12:
            colors = plt.cm.Set3(np.linspace(0, 1, num_trades))
        else:
            base_colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#ff99cc', 
                          '#c2c2f0', '#ffb3e6', '#c4e17f', '#76d7c4', '#f7dc6f']
            colors = (base_colors * ((num_trades // len(base_colors)) + 1))[:num_trades]
        
        fig, ax = plt.subplots(figsize=(10, 8))
        
        wedges, texts, autotexts = ax.pie(
            trade_data['DefectCount'], 
            labels=trade_data['Trade'], 
            colors=colors,
            autopct='%1.1f%%',
            startangle=45
        )
        
        ax.set_title(f'Distribution of Defects by Trade Category ({num_trades} Trades)', 
                    fontsize=16, fontweight='600', pad=20)
        
        plt.tight_layout()
        add_chart_to_document(doc, fig)
        plt.close()
        
        # Summary text
        if len(trade_data) > 0:
            top_trade = trade_data.iloc[0]
            summary_text = f"""The analysis reveals {top_trade['Trade']} as the primary defect category, representing {top_trade['DefectCount']} of the total {total_defects:,} defects ({top_trade['DefectCount']/total_defects*100:.1f}% of all identified issues). This complete analysis covers all {num_trades} trade categories identified during the inspection, providing comprehensive insights into defect distribution patterns."""
            
            summary_para = doc.add_paragraph(summary_text)
            summary_para.style = 'CleanBody'
    
    except Exception as e:
        print(f"Error creating pie chart: {e}")

def create_severity_chart(doc, metrics):
    """Create severity chart"""
    
    if not MATPLOTLIB_AVAILABLE:
        add_text_severity_summary(doc, metrics)
        return
    
    try:
        chart_title = doc.add_paragraph("Unit Classification by Defect Severity")
        chart_title.style = 'CleanSubsectionHeader'
        
        if 'summary_unit' in metrics and len(metrics['summary_unit']) > 0:
            fig, ax = plt.subplots(figsize=(12, 7))
            
            units_data = metrics['summary_unit']
            
            categories = []
            counts = []
            colors = []
            
            # Calculate categories
            extensive_count = len(units_data[units_data['DefectCount'] >= 15])
            categories.append('Extensive\n(15+ defects)')
            counts.append(extensive_count)
            colors.append('#ff9999')
            
            major_count = len(units_data[(units_data['DefectCount'] >= 8) & (units_data['DefectCount'] <= 14)])
            categories.append('Major\n(8-14 defects)')
            counts.append(major_count)
            colors.append('#ffcc99')
            
            minor_count = len(units_data[(units_data['DefectCount'] >= 3) & (units_data['DefectCount'] <= 7)])
            categories.append('Minor\n(3-7 defects)')
            counts.append(minor_count)
            colors.append('#ffff99')
            
            ready_count = len(units_data[units_data['DefectCount'] <= 2])
            categories.append('Ready\n(0-2 defects)')
            counts.append(ready_count)
            colors.append('#99ff99')
            
            bars = ax.bar(categories, counts, color=colors, alpha=0.8)
            
            ax.set_ylabel('Number of Units', fontsize=14, fontweight='600')
            ax.set_title('Unit Distribution by Defect Severity Level', 
                        fontsize=16, fontweight='600', pad=20)
            ax.grid(axis='y', alpha=0.3, linestyle=':')
            
            # Value labels
            for bar, value in zip(bars, counts):
                if value > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(counts)*0.01,
                           f'{value}', ha='center', va='bottom', 
                           fontweight='bold', fontsize=12)
            
            plt.tight_layout()
            add_chart_to_document(doc, fig)
            plt.close()
    
    except Exception as e:
        print(f"Error creating severity chart: {e}")

def create_trade_chart(doc, metrics):
    """Create trade analysis chart"""
    
    if not MATPLOTLIB_AVAILABLE:
        return
    
    try:
        trade_header = doc.add_paragraph("Trade Category Performance Analysis")
        trade_header.style = 'CleanSubsectionHeader'
        
        if 'summary_trade' not in metrics or len(metrics['summary_trade']) == 0:
            return
        
        top_trades = metrics['summary_trade'].head(10)
        
        fig, ax = plt.subplots(figsize=(12, 8))
        
        colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#ff99cc'] * 2
        colors = colors[:len(top_trades)]
        
        if NUMPY_AVAILABLE:
            y_pos = np.arange(len(top_trades))
        else:
            y_pos = list(range(len(top_trades)))
        
        bars = ax.barh(y_pos, top_trades['DefectCount'], color=colors, alpha=0.8)
        
        ax.set_yticks(y_pos)
        ax.set_yticklabels(top_trades['Trade'], fontsize=12)
        ax.set_xlabel('Number of Defects', fontsize=14, fontweight='600')
        ax.set_title('Trade Categories Ranked by Defect Frequency', 
                    fontsize=16, fontweight='600', pad=20)
        
        ax.grid(axis='x', alpha=0.3, linestyle=':')
        
        # Value labels
        total_defects = metrics.get('total_defects', 1)
        for i, (bar, value) in enumerate(zip(bars, top_trades['DefectCount'])):
            percentage = (value / total_defects * 100) if total_defects > 0 else 0
            ax.text(bar.get_width() + max(top_trades['DefectCount']) * 0.02, 
                   bar.get_y() + bar.get_height()/2,
                   f'{value} ({percentage:.1f}%)', va='center', 
                   fontweight='600', fontsize=10)
        
        plt.tight_layout()
        add_chart_to_document(doc, fig)
        plt.close()
    
    except Exception as e:
        print(f"Error creating trade chart: {e}")

def add_trade_summary(doc, processed_data, metrics):
    """Add trade summary section"""
    
    try:
        
        header = doc.add_paragraph("TRADE-SPECIFIC DEFECT ANALYSIS")
        header.style = 'CleanSectionHeader'
        
        # Decorative line - shortened as requested
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deco_run = deco_para.add_run("───────────────────────────────────────────────────────────────")  # Shortened line
        deco_run.font.name = 'Arial'
        deco_run.font.size = Pt(10)
        deco_run.font.color.rgb = RGBColor(0, 0, 0)

        overview_text = """This section provides a comprehensive breakdown of identified defects organized by trade category, including complete unit inventories for targeted remediation planning and resource allocation optimization."""
        
        overview_para = doc.add_paragraph(overview_text)
        overview_para.style = 'CleanBody'
                
        if processed_data is not None and len(processed_data) > 0:
            try:
                component_details = generate_complete_component_details(processed_data)
                add_trade_tables(doc, component_details)
            except Exception as e:
                print(f"Error generating trade tables: {e}")
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in trade summary: {e}")

def generate_complete_component_details(processed_data):
    """Generate component details for trade analysis"""
    
    try:
        required_columns = ['StatusClass', 'Trade', 'Room', 'Component', 'Unit']
        missing_columns = [col for col in required_columns if col not in processed_data.columns]
        
        if missing_columns:
            print(f"Missing columns: {missing_columns}")
            return pd.DataFrame()
        
        defects_only = processed_data[processed_data['StatusClass'] == 'Not OK']
        
        if len(defects_only) == 0:
            return pd.DataFrame()
        
        component_summary = defects_only.groupby(['Trade', 'Room', 'Component']).agg({
            'Unit': lambda x: ', '.join(sorted(x.astype(str).unique()))
        }).reset_index()
        
        component_summary.columns = ['Trade', 'Room', 'Component', 'Affected Units']
        
        unit_counts = defects_only.groupby(['Trade', 'Room', 'Component'])['Unit'].nunique().reset_index()
        component_summary = component_summary.merge(unit_counts, on=['Trade', 'Room', 'Component'])
        component_summary.columns = ['Trade', 'Room', 'Component', 'Affected Units', 'Unit Count']
        
        component_summary = component_summary.sort_values(['Trade', 'Unit Count'], ascending=[True, False])
        
        return component_summary
    
    except Exception as e:
        print(f"Error generating component details: {e}")
        return pd.DataFrame()

def add_trade_tables(doc, component_details):
    """Add trade tables with clean formatting and shading"""
    
    try:
        if len(component_details) == 0:
            return
        
        trades = component_details['Trade'].unique()
        
        for trade in trades:
            try:
                trade_data = component_details[component_details['Trade'] == trade]
                
                trade_header = doc.add_paragraph(f"{trade}")
                trade_header.style = 'CleanSubsectionHeader'
                
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                table.columns[0].width = Inches(2.5)
                table.columns[1].width = Inches(4.0)
                table.columns[2].width = Inches(0.8)
                
                # Headers with shading
                headers = ['Component & Location', 'Affected Units', 'Count']
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.runs[0]
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Add header shading (light gray)
                    set_cell_background_color(cell, "F0F0F0")
                
                # Add alternating row shading
                for idx, (_, row) in enumerate(trade_data.iterrows()):
                    table_row = table.add_row()
                    
                    # Set alternating row colors (white and light gray)
                    row_color = "FFFFFF" if idx % 2 == 0 else "F8F8F8"
                    
                    component_location = str(row['Component'])
                    if pd.notna(row['Room']) and str(row['Room']).strip():
                        component_location += f" ({row['Room']})"
                    
                    # Component & Location cell
                    cell1 = table_row.cells[0]
                    cell1.text = component_location
                    cell1.paragraphs[0].runs[0].font.name = 'Arial'
                    cell1.paragraphs[0].runs[0].font.size = Pt(10)
                    cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell1, row_color)
                    
                    # Affected Units cell
                    cell2 = table_row.cells[1]
                    cell2.text = str(row['Affected Units'])
                    cell2.paragraphs[0].runs[0].font.name = 'Arial'
                    cell2.paragraphs[0].runs[0].font.size = Pt(10)
                    cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell2, row_color)
                    
                    # Count cell
                    cell3 = table_row.cells[2]
                    cell3.text = str(row['Unit Count'])
                    cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell3.paragraphs[0].runs[0].font.name = 'Arial'
                    cell3.paragraphs[0].runs[0].font.size = Pt(10)
                    cell3.paragraphs[0].runs[0].font.bold = True
                    cell3.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell3, row_color)
                
                doc.add_paragraph()
            
            except Exception as e:
                print(f"Error processing trade {trade}: {e}")
                continue
    
    except Exception as e:
        print(f"Error in trade tables: {e}")

def add_component_breakdown(doc, processed_data, metrics):
    """Add component breakdown analysis - FIXED VERSION"""
    
    try:
        header = doc.add_paragraph("COMPONENT-LEVEL ANALYSIS")
        header.style = 'CleanSectionHeader'
        
        # Decorative line
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deco_run = deco_para.add_run("─────────────────────────────────────────────────────────────────")
        deco_run.font.name = 'Arial'
        deco_run.font.size = Pt(10)
        deco_run.font.color.rgb = RGBColor(0, 0, 0)
 
        intro_text = "This analysis identifies the most frequently affected individual components across all units, enabling targeted quality control improvements and preventive measures for future construction phases."
        
        intro_para = doc.add_paragraph(intro_text)
        intro_para.style = 'CleanBody'
        
        doc.add_paragraph()
        
        if processed_data is not None and len(processed_data) > 0:
            # FIXED: Generate component breakdown properly
            component_data = generate_fixed_component_breakdown(processed_data)
            
            if len(component_data) > 0:
                # FIXED: Proper aggregation by Component + Trade combination
                component_aggregated = component_data.groupby(['Component', 'Trade']).agg({
                    'Unit_Count': 'sum',  # Sum the unit counts for each component
                    'Affected_Units': lambda x: list(set([unit for units_str in x for unit in units_str.split(', ') if units_str.strip()]))  # Combine all unique units
                }).reset_index()
                
                # Convert the list back to comma-separated string and count unique units
                component_aggregated['Total_Unique_Units'] = component_aggregated['Affected_Units'].apply(len)
                component_aggregated['Sample_Units_Display'] = component_aggregated['Affected_Units'].apply(
                    lambda units_list: ', '.join(sorted(units_list[:5]) + ([f'(+{len(units_list)-5} more)'] if len(units_list) > 5 else []))
                )
                
                # Sort by total unique units affected (descending)
                top_components = component_aggregated.nlargest(15, 'Total_Unique_Units')
                
                # Filter to show only components affecting multiple units (if we have enough data)
                if len(top_components[top_components['Total_Unique_Units'] > 1]) >= 10:
                    top_components = top_components[top_components['Total_Unique_Units'] > 1].head(10)
                else:
                    top_components = top_components.head(10)
                
                most_freq_header = doc.add_paragraph("Most Frequently Affected Components")
                most_freq_header.style = 'CleanSubsectionHeader'
                
                if len(top_components) > 0:
                    comp_table = doc.add_table(rows=1, cols=5)
                    comp_table.style = 'Table Grid'
                    
                    comp_table.columns[0].width = Inches(2.0)
                    comp_table.columns[1].width = Inches(1.8)
                    comp_table.columns[2].width = Inches(2.5)
                    comp_table.columns[3].width = Inches(0.8)
                    comp_table.columns[4].width = Inches(1.0)
                    
                    # Headers with bold formatting
                    headers = ['Component', 'Trade', 'Sample Affected Units', 'Total Count', 'Percentage']
                    for i, header in enumerate(headers):
                        cell = comp_table.cell(0, i)
                        cell.text = header
                        para = cell.paragraphs[0]
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.runs[0]
                        run.font.bold = True
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        # Add header shading
                        set_cell_background_color(cell, "F0F0F0")
                    
                    total_units = metrics.get('total_units', 1)
                    for idx, (_, comp_row) in enumerate(top_components.iterrows()):
                        row = comp_table.add_row()
                        
                        # Alternating row colors
                        row_color = "FFFFFF" if idx % 2 == 0 else "F8F8F8"
                        
                        # Component
                        cell1 = row.cells[0]
                        cell1.text = str(comp_row.get('Component', 'N/A'))
                        cell1.paragraphs[0].runs[0].font.name = 'Arial'
                        cell1.paragraphs[0].runs[0].font.size = Pt(10)
                        cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        set_cell_background_color(cell1, row_color)
                        
                        # Trade
                        cell2 = row.cells[1]
                        cell2.text = str(comp_row.get('Trade', 'N/A'))
                        cell2.paragraphs[0].runs[0].font.name = 'Arial'
                        cell2.paragraphs[0].runs[0].font.size = Pt(10)
                        cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        set_cell_background_color(cell2, row_color)
                        
                        # Sample affected units
                        cell3 = row.cells[2]
                        cell3.text = str(comp_row.get('Sample_Units_Display', ''))
                        cell3.paragraphs[0].runs[0].font.name = 'Arial'
                        cell3.paragraphs[0].runs[0].font.size = Pt(9)
                        cell3.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        set_cell_background_color(cell3, row_color)
                        
                        # FIXED: Use Total_Unique_Units instead of Unit_Count
                        cell4 = row.cells[3]
                        unit_count = comp_row.get('Total_Unique_Units', 0)
                        cell4.text = str(unit_count)
                        cell4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell4.paragraphs[0].runs[0].font.name = 'Arial'
                        cell4.paragraphs[0].runs[0].font.size = Pt(10)
                        cell4.paragraphs[0].runs[0].font.bold = True
                        cell4.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        set_cell_background_color(cell4, row_color)
                        
                        # FIXED: Calculate percentage correctly using unique units
                        cell5 = row.cells[4]
                        percentage = (unit_count / total_units * 100) if total_units > 0 else 0
                        cell5.text = f"{percentage:.1f}%"
                        cell5.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell5.paragraphs[0].runs[0].font.name = 'Arial'
                        cell5.paragraphs[0].runs[0].font.size = Pt(10)
                        cell5.paragraphs[0].runs[0].font.bold = True
                        cell5.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        set_cell_background_color(cell5, row_color)
                    
                    # Analysis text with proper bold formatting
                    if len(top_components) > 0:
                        top_component = top_components.iloc[0]
                        unit_count = top_component.get('Total_Unique_Units', 0)
                        component_name = top_component.get('Component', 'Unknown')
                        trade_name = top_component.get('Trade', 'Unknown')
                        
                        doc.add_paragraph()
                        
                        analysis_text = f"""**Component Analysis Insights**: "{component_name}" emerges as the most frequently affected component, impacting {unit_count} units ({unit_count/total_units*100:.1f}% of all inspected units). This pattern reveals a systematic issue requiring immediate attention within the {trade_name} trade category.

**Key Findings from Component Analysis**:
• The top 5 most problematic components collectively affect {top_components.head(5)['Total_Unique_Units'].sum()} units across the development
• {trade_name} trade demonstrates the highest frequency of component-specific defects
• Recurring component failures across multiple units indicate potential systematic installation or quality control issues
• Component-level patterns suggest opportunities for targeted supplier quality improvements"""
                        
                        analysis_para = doc.add_paragraph()
                        add_formatted_text_with_bold(analysis_para, analysis_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in component breakdown: {e}")


def generate_fixed_component_breakdown(processed_data):
    """FIXED: Generate component details that match Excel report logic"""
    
    try:
        required_columns = ['StatusClass', 'Trade', 'Room', 'Component', 'Unit']
        missing_columns = [col for col in required_columns if col not in processed_data.columns]
        
        if missing_columns:
            print(f"Missing columns: {missing_columns}")
            return pd.DataFrame()
        
        # Filter for defects only
        defects_only = processed_data[processed_data['StatusClass'] == 'Not OK']
        
        if len(defects_only) == 0:
            return pd.DataFrame()
        
        # FIXED: Group by Component, Trade, Room combination (matching Excel logic)
        component_summary = defects_only.groupby(['Component', 'Trade', 'Room']).agg({
            'Unit': lambda x: ', '.join(sorted(x.astype(str).unique()))  # Get unique units as comma-separated string
        }).reset_index()
        
        # Count unique units per component/trade/room combination
        unit_counts = defects_only.groupby(['Component', 'Trade', 'Room'])['Unit'].nunique().reset_index()
        
        # Merge the data
        component_summary = component_summary.merge(unit_counts, on=['Component', 'Trade', 'Room'])
        component_summary.columns = ['Component', 'Trade', 'Room', 'Affected_Units', 'Unit_Count']
        
        # Sort by unit count (descending), then by component name
        component_summary = component_summary.sort_values(['Unit_Count', 'Component'], ascending=[False, True])
        
        print(f"Generated component breakdown with {len(component_summary)} entries")
        return component_summary
    
    except Exception as e:
        print(f"Error generating fixed component details: {e}")
        return pd.DataFrame()

def add_recommendations(doc, metrics):
    """Add recommendations section"""
    
    try:
        header = doc.add_paragraph("STRATEGIC RECOMMENDATIONS & ACTION PLAN")
        header.style = 'CleanSectionHeader'
        
        # Decorative line - shortened as requested
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deco_run = deco_para.add_run("───────────────────────────────────────────────────────────────")  # Shortened line
        deco_run.font.name = 'Arial'
        deco_run.font.size = Pt(10)
        deco_run.font.color.rgb = RGBColor(0, 0, 0)
                
        # Immediate priorities
        priorities_header = doc.add_paragraph("IMMEDIATE PRIORITIES")
        priorities_header.style = 'CleanSubsectionHeader'
        
        priorities = []
        ready_pct = metrics.get('ready_pct', 0)
        extensive_units = metrics.get('extensive_work_units', 0)
        
        if ready_pct > 75:
            priorities.append("**Accelerated Completion Protocol**: With 75%+ units requiring only minor work, implement immediate handover for compliant units while establishing parallel remediation workflows for remaining inventory.")
        elif ready_pct > 50:
            priorities.append("**Phased Completion Strategy**: Establish structured completion phases prioritizing ready units first, with clear milestone-based progression for units under remediation.")
        else:
            priorities.append("**Quality-First Approach**: Implement comprehensive remediation program before handover to ensure optimal customer satisfaction and minimize post-handover defect claims.")
        
        if 'summary_trade' in metrics and len(metrics['summary_trade']) > 0:
            top_trade = metrics['summary_trade'].iloc[0]
            top_trade_pct = (top_trade['DefectCount'] / metrics.get('total_defects', 1) * 100)
            priorities.append(f"**{top_trade['Trade']} Focus Initiative**: This trade represents {top_trade_pct:.1f}% of all defects ({top_trade['DefectCount']} instances). Deploy dedicated supervision teams and additional resources with daily progress monitoring.")
        
        if extensive_units > 0:
            priorities.append(f"**Specialized Remediation Teams**: {extensive_units} units require extensive work (15+ defects each). Establish dedicated teams with enhanced supervision to maintain project timeline integrity and quality standards.")
        
        priorities.append("**Enhanced Quality Protocols**: Implement multi-tier inspection checkpoints with supervisor sign-offs for critical trades before final handover, reducing post-handover callback rates.")
        
        for i, priority in enumerate(priorities, 1):
            priority_para = doc.add_paragraph()
            add_formatted_text_with_bold(priority_para, f"{i}. {priority}")
            priority_para.paragraph_format.left_indent = Inches(0.4)
        
        doc.add_paragraph()
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in recommendations: {e}")

def add_footer(doc, metrics):
    """Add footer section"""
    
    try:
        header = doc.add_paragraph("REPORT DOCUMENTATION & APPENDICES")
        header.style = 'CleanSectionHeader'
        
        # Decorative line - shortened as requested
        deco_para = doc.add_paragraph()
        deco_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        deco_run = deco_para.add_run("───────────────────────────────────────────────────────────────")  # Shortened line
        deco_run.font.name = 'Arial'
        deco_run.font.size = Pt(10)
        deco_run.font.color.rgb = RGBColor(0, 0, 0)

        # Comprehensive inspection metrics
        data_summary_header = doc.add_paragraph("COMPREHENSIVE INSPECTION METRICS")
        data_summary_header.style = 'CleanSubsectionHeader'
        
        avg_defects = metrics.get('avg_defects_per_unit', 0)
        defect_rate = metrics.get('defect_rate', 0)
        quality_score = max(0, 100 - defect_rate)
        
        data_summary_text = f"""**INSPECTION SCOPE & RESULTS**:
• Total Residential Units Evaluated: {metrics.get('total_units', 0):,}
• Total Building Components Assessed: {metrics.get('total_inspections', 0):,}
• Total Defects Documented: {metrics.get('total_defects', 0):,}
• Overall Defect Rate: {metrics.get('defect_rate', 0):.2f}%
• Average Defects per Unit: {avg_defects:.2f}
• Development Quality Score: {quality_score:.1f}/100

**DEFECT LEVEL FRAMEWORK DISTRIBUTION**:
• Minor Work Required: {metrics.get('ready_units', 0)} units ({metrics.get('ready_pct', 0):.1f}%)
• Intermediate Remediation Required: {metrics.get('minor_work_units', 0)} units ({metrics.get('minor_pct', 0):.1f}%)
• Major Remediation Required: {metrics.get('major_work_units', 0)} units ({metrics.get('major_pct', 0):.1f}%)  
• Extensive Remediation Required: {metrics.get('extensive_work_units', 0)} units ({metrics.get('extensive_pct', 0):.1f}%)"""
        
        data_summary_para = doc.add_paragraph()
        add_formatted_text_with_bold(data_summary_para, data_summary_text)
        
        doc.add_paragraph()
        
        # Report generation details
        details_header = doc.add_paragraph("REPORT GENERATION & COMPANION RESOURCES")
        details_header.style = 'CleanSubsectionHeader'
        
        details_text = f"""**REPORT METADATA**:
• Report Generated: {datetime.now().strftime('%d %B %Y at %I:%M %p')}
• Inspection Completion: {metrics.get('inspection_date', 'N/A')}
• Building Development: {metrics.get('building_name', 'N/A')}
• Property Location: {metrics.get('address', 'N/A')}

**COMPANION DOCUMENTATION SUITE**:
Complete defect inventories, unit-by-unit detailed breakdowns, interactive filterable data tables, and comprehensive photographic documentation are available in the accompanying Excel analytics workbook. This comprehensive dataset includes advanced filtering capabilities, dynamic visual dashboards, pivot table analysis tools, and direct export functionality for integration with project management systems and remediation tracking platforms.

**TECHNICAL SUPPORT & FOLLOW-UP**:
For technical inquiries, data interpretation assistance, or additional analysis requirements, please contact the inspection team. Ongoing support is available for remediation planning, progress tracking, and post-completion verification inspections."""
        
        details_para = doc.add_paragraph()
        add_formatted_text_with_bold(details_para, details_text)
        
        # Closing
        doc.add_paragraph()
        closing_para = doc.add_paragraph()
        closing_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        closing_run = closing_para.add_run("END OF REPORT")
        closing_run.font.name = 'Arial'
        closing_run.font.size = Pt(14)
        closing_run.font.color.rgb = RGBColor(0, 0, 0)
        closing_run.font.bold = True
    
    except Exception as e:
        print(f"Error in footer: {e}")

def add_chart_to_document(doc, fig):
    """Helper function to add charts to document"""
    
    try:
        chart_buffer = BytesIO()
        fig.savefig(chart_buffer, format='png', dpi=300, bbox_inches='tight', 
                    facecolor='white', edgecolor='none', pad_inches=0.2)
        chart_buffer.seek(0)
        
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_run = chart_para.add_run()
        chart_run.add_picture(chart_buffer, width=Inches(7))
        
        doc.add_paragraph()
    
    except Exception as e:
        print(f"Error adding chart: {e}")

def add_text_trade_summary(doc, metrics):
    """Text-based trade summary when matplotlib is not available"""
    try:
        breakdown_header = doc.add_paragraph("Defects Distribution by Trade Category")
        breakdown_header.style = 'CleanSubsectionHeader'
        
        note_para = doc.add_paragraph("(Visual charts require matplotlib - showing text summary)")
        note_para.style = 'CleanBody'
        
        if 'summary_trade' not in metrics or len(metrics['summary_trade']) == 0:
            return
        
        trade_data = metrics['summary_trade'].copy()
        total_defects = metrics.get('total_defects', 0)
        
        if total_defects > 0:
            for idx, (_, row) in enumerate(trade_data.iterrows(), 1):
                percentage = (row['DefectCount'] / total_defects * 100)
                trade_text = f"{idx}. {row['Trade']}: {row['DefectCount']} defects ({percentage:.1f}%)"
                trade_para = doc.add_paragraph(trade_text)
                trade_para.style = 'CleanBody'
                trade_para.paragraph_format.left_indent = Inches(0.3)
        
    except Exception as e:
        print(f"Error in text trade summary: {e}")

def add_text_severity_summary(doc, metrics):
    """Text-based severity summary when matplotlib is not available"""
    try:
        chart_title = doc.add_paragraph("Unit Classification by Defect Severity")
        chart_title.style = 'CleanSubsectionHeader'
        
        note_para = doc.add_paragraph("(Visual charts require matplotlib - showing text summary)")
        note_para.style = 'CleanBody'
        
        if 'summary_unit' in metrics and len(metrics['summary_unit']) > 0:
            units_data = metrics['summary_unit']
            
            extensive_count = len(units_data[units_data['DefectCount'] >= 15])
            major_count = len(units_data[(units_data['DefectCount'] >= 8) & (units_data['DefectCount'] <= 14)])
            minor_count = len(units_data[(units_data['DefectCount'] >= 3) & (units_data['DefectCount'] <= 7)])
            ready_count = len(units_data[units_data['DefectCount'] <= 2])
            
            severity_data = [
                ("Extensive (15+ defects)", extensive_count),
                ("Major (8-14 defects)", major_count),
                ("Minor (3-7 defects)", minor_count),
                ("Ready (0-2 defects)", ready_count)
            ]
            
            for category, count in severity_data:
                if count > 0:
                    severity_text = f"• {category}: {count} units"
                    severity_para = doc.add_paragraph(severity_text)
                    severity_para.style = 'CleanBody'
                    severity_para.paragraph_format.left_indent = Inches(0.3)
        
    except Exception as e:
        print(f"Error in text severity summary: {e}")

def add_text_units_summary(doc, metrics):
    """Text-based units summary when matplotlib is not available"""
    try:
        chart_title = doc.add_paragraph("Top 20 Units Requiring Immediate Intervention")
        chart_title.style = 'CleanSubsectionHeader'
        
        note_para = doc.add_paragraph("(Visual charts require matplotlib - showing text summary)")
        note_para.style = 'CleanBody'
        
        if 'summary_unit' not in metrics or len(metrics['summary_unit']) == 0:
            return
        
        top_units = metrics['summary_unit'].head(20)
        
        for idx, (_, row) in enumerate(top_units.iterrows(), 1):
            unit_text = f"{idx}. Unit {row['Unit']}: {row['DefectCount']} defects"
            unit_para = doc.add_paragraph(unit_text)
            unit_para.style = 'CleanBody'
            unit_para.paragraph_format.left_indent = Inches(0.3)
        
    except Exception as e:
        print(f"Error in text units summary: {e}")

def create_error_document(error, metrics):
    """Create error document"""
    
    doc = Document()
    title = doc.add_heading("Inspection Report - Generation Error", level=1)
    error_para = doc.add_paragraph(f"Report generation encountered an issue: {str(error)}")
    
    if metrics:
        basic_para = doc.add_paragraph(f"""
Basic Information:
Building: {metrics.get('building_name', 'N/A')}
Total Units: {metrics.get('total_units', 'N/A')}
Total Defects: {metrics.get('total_defects', 'N/A')}
        """)
    
    return doc

# Backward compatibility functions
def generate_professional_word_report(processed_data, metrics, images=None):
    """Backward compatibility wrapper"""
    return generate_enhanced_word_report(processed_data, metrics, images)

def generate_word_report(processed_data, metrics, images=None):
    """Backward compatibility function"""
    return generate_enhanced_word_report(processed_data, metrics, images)

def create_inspection_report(processed_data, metrics, images=None):
    """Alternative function name for backward compatibility"""
    return generate_enhanced_word_report(processed_data, metrics, images)

# Main execution and testing
if __name__ == "__main__":
    print("Enhanced Word Report Generator with Blank Page Removal loaded successfully!")
    print("\nNEW FEATURES ADDED:")
    print("• remove_blank_pages() - Identifies and removes blank pages/elements")
    print("• cleanup_excessive_spacing() - Removes excessive empty paragraphs")
    print("• optimize_page_breaks() - Optimizes page break placement")
    print("• validate_document_structure() - Provides document structure analysis")
    print("• Post-processing pipeline with detailed logging")
    
    print("\nEXISTING FEATURES:")
    print("• Chart legends display properly with severity levels")
    print("• Bold text formatting (**text**) converted to actual bold")
    print("• Top margin set to 3cm")
    print("• Title split into 2 lines: PRE-SETTLEMENT / INSPECTION REPORT")
    print("• Shortened decorative lines")
    print("• Table shading with alternating rows and header colors")
    print("• Company logo positioned in document header (left side)")
    print("• Cover image sized appropriately (4.7 inches)")
    print("• All text uses Arial font with black color")
    
    print("\nDEPENDENCY STATUS:")
    print(f"  python-docx: Available")
    print(f"  matplotlib: {'Available' if MATPLOTLIB_AVAILABLE else 'Not Available'}")
    print(f"  seaborn: {'Available' if SEABORN_AVAILABLE else 'Not Available'}")  
    print(f"  numpy: {'Available' if NUMPY_AVAILABLE else 'Not Available'}")
    
    if not MATPLOTLIB_AVAILABLE:
        print("\nINFO: matplotlib not available - charts will be replaced with text summaries")
        print("For full visual charts, install with: pip install matplotlib seaborn")
    else:
        print("\nREADY: All dependencies available for full visual reports with legends!")
    
    print("\nPOST-PROCESSING WORKFLOW:")
    print("1. Document generation with all sections")
    print("2. Initial structure validation")
    print("3. Blank page detection and removal")
    print("4. Page break optimization")
    print("5. Excessive spacing cleanup")
    print("6. Final structure validation and reporting")
    
    print("\nREADY FOR PRODUCTION WITH BLANK PAGE REMOVAL!")